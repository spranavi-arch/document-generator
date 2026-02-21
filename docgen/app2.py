"""
Streamlit UI for docgen: upload two samples, optional CURL, run pipeline step-by-step.
Live UI: see sections → prompts → field values → draft building section by section → final draft + download.
Uses .env from backend folder.
"""
import re
import sys
import importlib.util
from pathlib import Path
from io import BytesIO

_root = Path(__file__).resolve().parent.parent
_docgen = Path(__file__).resolve().parent
for p in (str(_root), str(_docgen)):
    if p not in sys.path:
        sys.path.insert(0, p)

import streamlit as st

# Optional rich text editors (Word-like toolbar)
try:
    from streamlit_quill import st_quill
    _HAS_QUILL = True
except ImportError:
    _HAS_QUILL = False
try:
    from streamlit_lexical import streamlit_lexical
    _HAS_LEXICAL = True
except ImportError:
    _HAS_LEXICAL = False

# Plain text <-> simple HTML for rich editor (matches formatting/utils/html_to_docx markers)
_SECTION_UNDERLINE_MARKER = "[SECTION_UNDERLINE]"


def _plain_text_to_simple_html(text: str) -> str:
    """Wrap plain text in simple HTML for the editor. [SECTION_UNDERLINE] -> <hr class=\"section-underline\">."""
    if not text:
        return "<p><br></p>"
    parts = []
    for para in (text or "").split("\n\n"):
        para = (para or "").strip()
        if para == _SECTION_UNDERLINE_MARKER:
            parts.append('<hr class="section-underline">')
        else:
            para = (para or "").replace("\n", "<br>")
            if para:
                parts.append(f"<p>{para}</p>")
            else:
                parts.append("<p><br></p>")
    return "".join(parts) if parts else "<p><br></p>"


def _simple_html_to_plain_text(html: str) -> str:
    """Extract plain text from simple HTML. <hr class=\"section-underline\"> -> [SECTION_UNDERLINE]. Strips other tags."""
    if not html:
        return ""
    text = re.sub(
        r'<hr[^>]*class="[^"]*section-underline[^"]*"[^>]*>',
        "\n\n" + _SECTION_UNDERLINE_MARKER + "\n\n",
        html,
        flags=re.I,
    )
    text = re.sub(r"<hr[^>]*>", "\n\n", text, flags=re.I)
    text = re.sub(r"</p>\s*<p>", "\n\n", text)
    text = text.replace("<p>", "").replace("</p>", "").replace("<br>", "\n").replace("<br/>", "\n")
    text = re.sub(r"<[^>]+>", "", text)  # strip remaining tags (e.g. <strong>, <em>)
    return text.strip()

from docgen.sectioner import divide_into_sections
from docgen.extractor import extract_sections_from_docs
from docgen.section_prompt_generator import generate_prompt_and_fields
from docgen.field_fetcher import (
    call_chat_api_with_question_debug,
    fetch_all_fields_via_chat,
    fetch_case_summary,
    _default_question_for_field,
)
from docgen.question_generator import generate_questions_for_fields
from docgen.section_generator import generate_section
from docgen.assembler import assemble
from docgen.section_formatting_prompt_generator import generate_section_formatting_instructions

# -----------------------------------------------------------------------------
# Formatting integration (code in formatting/)
# -----------------------------------------------------------------------------
def _load_formatting_backend():
    """Load formatting backend module (sys.path and imports isolated). Returns the module."""
    fmt_dir = _root / "formatting"
    fmt_dir_str = str(fmt_dir)
    backend_file = fmt_dir / "backend.py"
    saved_path = sys.path.copy()
    saved_modules = {}
    for key in list(sys.modules.keys()):
        if key == "utils" or key.startswith("utils.") or key == "backend":
            saved_modules[key] = sys.modules.pop(key)
    sys.path = [fmt_dir_str]
    try:
        spec = importlib.util.spec_from_file_location("formatting_backend", backend_file)
        fmt_backend = importlib.util.module_from_spec(spec)
        sys.modules["formatting_backend"] = fmt_backend
        spec.loader.exec_module(fmt_backend)
        return fmt_backend
    finally:
        sys.path = saved_path
        for key, mod in saved_modules.items():
            sys.modules[key] = mod
        sys.modules.pop("formatting_backend", None)


def _run_formatting_llm(final_draft: str, sample1_bytes: bytes) -> tuple[str, str, bytes, dict]:
    """Call formatting backend: use first sample as DOCX template, final_draft as text. Returns (output_path, preview_text, docx_bytes, prompt_preview)."""
    fmt_backend = _load_formatting_backend()
    template_file = BytesIO(sample1_bytes)
    output_path, preview_text = fmt_backend.process_document(final_draft, template_file)
    with open(output_path, "rb") as f:
        docx_bytes = f.read()
    return output_path, preview_text, docx_bytes, {}


def _run_formatting_section_by_section(
    sections_list: list,
    extracted_samples: list,
    draft_parts: list,
    sample1_bytes: bytes,
) -> tuple[str, str, bytes, list, list]:
    """Step 5a: Generate per-section formatting prompts. Step 5b: Format section by section. Returns (output_path, preview_text, docx_bytes, section_formatting_prompts, blocks)."""
    fmt_backend = _load_formatting_backend()
    template_file = BytesIO(sample1_bytes)
    schema = fmt_backend.get_schema_from_template(template_file)
    template_file.seek(0)
    section_formatting_prompts = generate_section_formatting_instructions(
        sections_list,
        extracted_samples,
        schema.get("template_content", []),
        schema.get("style_guide", ""),
    )
    output_path, preview_text, all_blocks = fmt_backend.process_document_section_by_section(
        draft_parts,
        section_formatting_prompts,
        BytesIO(sample1_bytes),
    )
    with open(output_path, "rb") as f:
        docx_bytes = f.read()
    return output_path, preview_text, docx_bytes, section_formatting_prompts, all_blocks


def _run_formatting_structural(
    sections_list: list,
    extracted_samples: list,
    draft_parts: list,
    sample1_bytes: bytes,
) -> tuple[str, str, bytes, list, list, dict]:
    """Run 7-step structural formatting. Returns (output_path, preview_text, docx_bytes, section_fmt_prompts, mapped_blocks, structural_result)."""
    fmt_backend = _load_formatting_backend()
    template_file = BytesIO(sample1_bytes)
    schema = fmt_backend.get_schema_from_template(template_file)
    template_file.seek(0)
    section_formatting_prompts = generate_section_formatting_instructions(
        sections_list,
        extracted_samples,
        schema.get("template_content", []),
        schema.get("style_guide", ""),
    )
    output_path, preview_text, mapped_blocks, structural_result = fmt_backend.process_document_structural(
        draft_parts,
        section_formatting_prompts,
        BytesIO(sample1_bytes),
    )
    with open(output_path, "rb") as f:
        docx_bytes = f.read()
    return output_path, preview_text, docx_bytes, section_formatting_prompts, mapped_blocks, structural_result


def _edited_preview_to_blocks(edited_text: str, stored_blocks: list) -> list:
    """Convert edited preview text back to blocks using stored block types. Preview uses \\n\\n between paragraphs; [SECTION_UNDERLINE] for section underlines. If stored_blocks is empty, infers blocks from segments ([SECTION_UNDERLINE] -> section_underline, else paragraph)."""
    segments = [s.strip() for s in edited_text.split("\n\n")]
    if not stored_blocks:
        new_blocks = []
        for seg in segments:
            if seg == "[SECTION_UNDERLINE]":
                new_blocks.append(("section_underline", ""))
            else:
                new_blocks.append(("paragraph", seg))
        return new_blocks
    new_blocks = []
    seg_idx = 0
    for block_type, _ in stored_blocks:
        if block_type == "page_break":
            new_blocks.append(("page_break", ""))
            continue
        if seg_idx >= len(segments):
            break
        seg = segments[seg_idx]
        seg_idx += 1
        if block_type == "section_underline":
            new_blocks.append(("section_underline", ""))
            continue
        new_blocks.append((block_type, seg))
    while seg_idx < len(segments):
        new_blocks.append(("paragraph", segments[seg_idx]))
        seg_idx += 1
    return new_blocks


def _blocks_to_draft_text(blocks: list) -> str:
    """Convert stored blocks to draft text (\\n\\n between paragraphs, [SECTION_UNDERLINE] for underlines) for process_document."""
    parts = []
    for kind, text in blocks:
        if kind == "page_break":
            continue
        if kind == "section_underline":
            parts.append(_SECTION_UNDERLINE_MARKER)
        else:
            parts.append((text or "").strip())
    return "\n\n".join(p for p in parts if p).strip()


def _run_editor_rebuild(blocks: list, sample1_bytes: bytes, formatting_overrides: dict) -> tuple[bytes, str]:
    """Rebuild DOCX from blocks by converting to draft text and calling formatting backend process_document. Returns (docx_bytes, preview_text)."""
    draft_text = _blocks_to_draft_text(blocks)
    fmt_backend = _load_formatting_backend()
    output_path, preview_text = fmt_backend.process_document(draft_text, BytesIO(sample1_bytes))
    with open(output_path, "rb") as f:
        return f.read(), preview_text


def _run_formatting_draft_driven(final_draft: str, sample1_bytes: bytes) -> tuple[str, str, bytes]:
    """Format using template DOCX styles; structure from draft. Uses formatting backend process_document (LLM + template). Returns (output_path, preview_text, docx_bytes)."""
    fmt_backend = _load_formatting_backend()
    output_path, preview_text = fmt_backend.process_document(final_draft, BytesIO(sample1_bytes))
    with open(output_path, "rb") as f:
        docx_bytes = f.read()
    return output_path, preview_text, docx_bytes


def _run_editor_rebuild_draft_driven(edited_text: str, sample1_bytes: bytes) -> tuple[bytes, str]:
    """Rebuild DOCX from edited text using formatting backend process_document. Returns (docx_bytes, preview_text)."""
    fmt_backend = _load_formatting_backend()
    output_path, preview_text = fmt_backend.process_document(edited_text, BytesIO(sample1_bytes))
    with open(output_path, "rb") as f:
        return f.read(), preview_text


# -----------------------------------------------------------------------------
# Config & helpers
# -----------------------------------------------------------------------------

st.set_page_config(page_title="DocGen – Section-based Draft", layout="wide")
st.title("Document Generator (Section-based)")

OLE_MAGIC = b"\xd0\xcf\x11\xe0"


def file_to_text(data: bytes, filename: str) -> str:
    """Extract plain text from uploaded file (.txt or .docx)."""
    if data.startswith(OLE_MAGIC):
        raise ValueError("Legacy .doc is not supported. Use .docx or .txt.")
    name = (filename or "").lower()
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("cp1252")


def text_to_docx_bytes(text: str) -> bytes:
    """Build a .docx in memory from plain text (paragraphs split on double newline)."""
    from docx import Document
    doc = Document()
    for block in (text or "").split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block.replace("\n", " "))
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# -----------------------------------------------------------------------------
# Sidebar: inputs
# -----------------------------------------------------------------------------

with st.sidebar:
    st.header("Inputs")
    sample1 = st.file_uploader("Sample document 1", type=["txt", "docx"], key="s1")
    sample2 = st.file_uploader("Sample document 2", type=["txt", "docx"], key="s2")
    st.markdown("---")
    testing_skip_api = st.checkbox(
        "Testing mode (skip API)",
        value=False,
        help="Generate draft without calling the API. Use this to test the formatting option only.",
    )
    st.markdown("---")
    st.subheader("API / CURL (for field data)")
    curl_input = st.text_area(
        "Paste CURL command (chat-with-case API)",
        height=120,
        placeholder="curl 'https://...' -H 'authorization: Bearer ...' --data-raw '{\"content\":\"...\"}'",
        help="Optional. We call this API per field for field data.",
    )
    extra_context = st.text_area(
        "Extra context / case summary",
        height=100,
        placeholder="Optional: case summary or extra data if not using API.",
    )
    st.caption("Uses backend/.env for Azure OpenAI.")


# -----------------------------------------------------------------------------
# Run pipeline step-by-step with live UI
# -----------------------------------------------------------------------------

def run_pipeline():
    if not sample1 or not sample2:
        st.error("Please upload both sample documents.")
        return
    try:
        data1 = sample1.read()
        sample1.seek(0)
        s1 = file_to_text(data1, sample1.name or "")
        sample1.seek(0)
        s2 = file_to_text(sample2.read(), sample2.name or "")
        sample2.seek(0)
        # Store first sample for formatting step (template must be DOCX)
        st.session_state["docgen_sample1_bytes"] = data1
        st.session_state["docgen_sample1_name"] = sample1.name or ""
    except Exception as e:
        st.error(f"Could not read files: {e}")
        return

    curl_str = (curl_input or "").strip() if not testing_skip_api else ""
    ctx = (extra_context or "").strip()
    if testing_skip_api:
        st.info("Testing mode: API calls skipped. Draft will use empty field data so you can test formatting.")

    # ----- Step 1: Identify sections (name + purpose only) -----
    st.subheader("Step 1: Identify sections")
    st.caption("Logical sections are identified from both documents. Full text is extracted in Step 2 in small chunks so nothing is truncated.")
    try:
        blueprint = divide_into_sections(s1, s2)
    except ValueError as e:
        st.error(str(e))
        return
    sections_list = blueprint["sections"]
    st.success(f"Identified {len(sections_list)} sections.")
    for i, sec in enumerate(sections_list, 1):
        st.markdown(f"**{i}. {sec.get('name', '?')}** — {sec.get('purpose', '') or '(no purpose)'}")

    # ----- Step 2: Extract section text (chunked so full doc content fits) then generate prompts -----
    st.subheader("Step 2: Extract section text and generate prompts")
    st.caption("Extracting text for each section in small chunks so the whole document is used without truncation. Then building prompts.")
    bar = st.progress(0, text="Extracting from document 1...")
    extracted_samples = extract_sections_from_docs(s1, s2, sections_list)
    bar.progress(0.5, text="Building prompts...")
    section_prompts_list = []
    for i, sec in enumerate(sections_list):
        bar.progress(0.5 + 0.5 * (i + 1) / len(sections_list), text=sec["name"])
        sample_text = extracted_samples[i] if i < len(extracted_samples) else ""
        section_prompts_list.append(
            generate_prompt_and_fields(sec["name"], sec.get("purpose", ""), sample_text)
        )
    bar.progress(1.0, text="Done.")
    st.success("Prompts and required fields ready.")
    with st.expander("View extracted sample text for each section (how extraction worked)"):
        st.caption("Section text was extracted in chunks (a few sections per call) so the full document content is used and nothing is truncated.")
        for i, sec in enumerate(sections_list):
            sample_text = extracted_samples[i] if i < len(extracted_samples) else ""
            with st.expander(f"{i + 1}. {sec['name']} — {len(sample_text)} chars"):
                if sample_text:
                    st.text_area(
                        "Extracted text",
                        value=sample_text,
                        height=min(400, max(120, 80 + sample_text.count("\n") * 18)),
                        key=f"extracted_sample_{i}",
                        disabled=True,
                        label_visibility="collapsed",
                    )
                else:
                    st.info("No text extracted for this section.")
    st.session_state["docgen_extracted_samples"] = extracted_samples
    with st.expander("View prompt and required fields for each section"):
        for i, sec in enumerate(sections_list):
            info = section_prompts_list[i]
            st.markdown(f"### {sec['name']}")
            st.caption("Required fields: " + ", ".join(info.get("required_fields", [])))
            st.text(info.get("prompt", "")[:700] + ("..." if len(info.get("prompt", "")) > 700 else ""))

    # ----- Step 3: Fetch field values via API -----
    all_required = []
    seen = set()
    for i in range(len(sections_list)):
        for f in section_prompts_list[i].get("required_fields", []):
            if f and f not in seen:
                seen.add(f)
                all_required.append(f)

    field_values = {}
    if curl_str:
        st.subheader("Step 3: Fetching data via API")
        if all_required:
            with st.status("Generating questions for each field...", state="running"):
                field_to_question = generate_questions_for_fields(all_required)
            st.success(f"Generated {len(field_to_question)} questions.")
            with st.expander("Field → question used for API"):
                for f, q in field_to_question.items():
                    st.markdown(f"**{f}** → \"{q}\"")

            # One test request so user can see why answers might be empty
            first_field = all_required[0]
            first_question = field_to_question.get(first_field) or _default_question_for_field(first_field)
            debug = call_chat_api_with_question_debug(curl_str, first_question)
            with st.expander("API test (first request)"):
                st.caption(f"Question: \"{first_question}\"")
                if debug.get("error"):
                    st.error(debug["error"])
                st.code(f"Status: {debug.get('status_code')} | Response keys: {debug.get('response_keys', [])}")
                if debug.get("extracted_preview"):
                    st.success(f"Extracted: {debug['extracted_preview']}")
                else:
                    st.warning("No text could be extracted from the response. Check that your API returns a body with one of: content, answer, message, text, or choices[0].message.content.")

            status_placeholder = st.empty()
            progress = st.progress(0, text="Preparing...")

            def on_field_start(field_name: str, index: int, total: int):
                status_placeholder.markdown(f"**Fetching field {index} of {total}:** `{field_name}`")
                progress.progress(index / total, text=f"Fetching: {field_name} ({index}/{total})")

            field_values = fetch_all_fields_via_chat(
                curl_str, all_required, field_to_question, on_field_start=on_field_start
            )
            status_placeholder.markdown("**Done.** All fields fetched.")
            progress.progress(1.0, text="Done.")

            # Extra question: case summary (passed to every section)
            status_placeholder.markdown("**Fetching case summary…**")
            case_summary = fetch_case_summary(curl_str)
            field_values["case_summary"] = case_summary or ""
            if case_summary:
                st.caption("Case summary received and will be passed to each section.")
            status_placeholder.markdown("**Done.** All fields fetched.")

        with st.expander("Field → value (from API)"):
            for f, v in field_values.items():
                vstr = str(v)
                st.markdown(f"**{f}** → {vstr[:200] + '...' if len(vstr) > 200 else vstr or '(empty)'}")
    else:
        st.subheader("Step 3: Field data")
        st.info("No CURL provided. Using extra context only if provided.")

    if ctx:
        field_values["case_summary_or_context"] = ctx

    # ----- Step 4: Generate sections one by one (draft builds live) -----
    st.subheader("Step 4: Drafting section by section")
    st.caption("Draft updates below as each section is generated.")
    draft_placeholder = st.empty()
    draft_parts = []
    for i, sec in enumerate(sections_list):
        name = sec["name"]
        info = section_prompts_list[i]
        prompt = info.get("prompt", "")
        required_fields = info.get("required_fields", [])
        section_field_values = {f: field_values.get(f, "") for f in required_fields}
        if ctx:
            section_field_values["case_summary_or_context"] = ctx
        # Pass case summary to every section (from extra API question); use if section needs context
        case_summary = field_values.get("case_summary", "")
        if case_summary:
            section_field_values["case_summary"] = case_summary

        with st.status(f"Generating section {i + 1}/{len(sections_list)}: **{name}**", state="running"):
            text = generate_section(
                prompt, section_field_values,
                sample_text=extracted_samples[i] if i < len(extracted_samples) else "",
                section_name=name,
            )
            draft_parts.append(text)
        # Show growing draft after each section (unique key per iteration to avoid duplicate key error)
        current_draft = "\n\n".join(draft_parts)
        draft_placeholder.text_area(
            "Draft so far",
            value=current_draft,
            height=300,
            label_visibility="collapsed",
            key=f"draft_so_far_{i}",
        )

    # ----- Step 5: Assemble final draft -----
    # Pass ordered list so each section appears once (avoids repetition when blueprint has duplicate section names)
    final_draft = assemble(blueprint, draft_parts)

    st.session_state["docgen_final_draft"] = final_draft
    st.session_state["docgen_blueprint"] = blueprint
    st.session_state["docgen_sections_list"] = sections_list
    st.session_state["docgen_section_prompts"] = section_prompts_list
    st.session_state["docgen_draft_parts"] = draft_parts
    st.session_state["docgen_field_values"] = field_values

    st.success("Draft complete.")
    st.subheader("Final document draft")
    st.text_area(
        "Final draft",
        value=final_draft,
        height=420,
        label_visibility="collapsed",
        key="final_draft_area",
    )
    sample1_name = st.session_state.get("docgen_sample1_name", "")
    use_formatting = sample1_name and sample1_name.lower().endswith(".docx")
    if use_formatting:
        st.subheader("Step 5: Format document (draft-driven)")
        st.caption("Structure comes from your final draft (paragraph breaks, headings, lists). Styling (fonts, spacing, alignment) comes from the sample DOCX. No slot mapping — the formatted doc matches the draft’s structure.")
        with st.spinner("Applying sample styles to draft structure…"):
            try:
                out_path, preview, docx_bytes = _run_formatting_draft_driven(
                    final_draft,
                    st.session_state["docgen_sample1_bytes"],
                )
                st.session_state["docgen_formatted_preview"] = preview
                st.session_state["docgen_formatted_docx_bytes"] = docx_bytes
                st.session_state["docgen_formatted_output_path"] = out_path
                st.session_state["docgen_used_draft_driven"] = True
                st.session_state["docgen_section_formatting_prompts"] = None
                st.session_state["docgen_formatted_blocks"] = None
                st.session_state["docgen_structural_result"] = None
                st.session_state["docgen_formatting_prompts"] = None
            except Exception as e:
                st.error(f"Formatting failed: {e}")
        if st.session_state.get("docgen_formatted_preview") is not None:
            st.success("Formatting complete. Formatted document has the same structure as your final draft.")
            st.subheader("Formatted document")
            st.text_area(
                "Formatted preview",
                value=st.session_state["docgen_formatted_preview"],
                height=280,
                label_visibility="collapsed",
                key="formatted_preview_run",
            )
            # Editor: change content or formatting before download
            with st.expander("Edit before download", expanded=True):
                st.caption("Edit the content below and/or set formatting options, then click **Apply edits and download**. For full Word-style editing, download the file and open it in [Word for the Web](https://word.cloud.microsoft) or desktop Word.")
                preview_text = st.session_state["docgen_formatted_preview"]
                initial_html = _plain_text_to_simple_html(preview_text)
                editor_plain = preview_text  # fallback for Apply
                if _HAS_QUILL:
                    st.caption("Use the toolbar for **bold**, *italic*, underline, alignment, lists.")
                    quill_content = st_quill(
                        value=initial_html,
                        html=True,
                        key="docgen_quill_editor_run",
                        toolbar=[
                            ["bold", "italic", "underline", "strike"],
                            [{"align": []}],
                            [{"list": "ordered"}, {"list": "bullet"}],
                            ["clean"],
                        ],
                    )
                    if quill_content is not None:
                        st.session_state["docgen_editor_content_html"] = quill_content
                    editor_plain = _simple_html_to_plain_text(
                        st.session_state.get("docgen_editor_content_html") or initial_html
                    )
                elif _HAS_LEXICAL:
                    st.caption("Edit below (plain text). Use **bold** / *italic* if your Lexical build supports it.")
                    lexical_content = streamlit_lexical(
                        value=preview_text,
                        placeholder="Edit document content…",
                        height=280,
                        debounce=500,
                        key="docgen_lexical_editor_run",
                    )
                    if lexical_content is not None:
                        st.session_state["docgen_editor_content_plain"] = lexical_content
                    editor_plain = st.session_state.get("docgen_editor_content_plain") or preview_text
                else:
                    st.caption("Install **streamlit-quill** for a toolbar editor (bold, italic, lists). Or download and open in [Word for the Web](https://word.cloud.microsoft).")
                    editor_preview = st.text_area(
                        "Content (edit as needed)",
                        value=preview_text,
                        height=260,
                        key="editor_preview_run",
                        label_visibility="collapsed",
                    )
                    editor_plain = st.session_state.get("editor_preview_run", editor_preview)
                col1, col2 = st.columns(2)
                with col1:
                    space_after_pt = st.number_input(
                        "Space after paragraph (pt); 0 = keep template",
                        min_value=0,
                        max_value=72,
                        value=0,
                        step=1,
                        key="editor_space_after_run",
                    )
                with col2:
                    body_font_size_pt = st.number_input(
                        "Body font size (pt); 0 = keep template",
                        min_value=0,
                        max_value=24,
                        value=0,
                        step=1,
                        key="editor_font_size_run",
                    )
                overrides = {}
                if space_after_pt > 0:
                    overrides["space_after_pt"] = space_after_pt
                if body_font_size_pt > 0:
                    overrides["body_font_size_pt"] = body_font_size_pt
                if st.button("Apply edits and download", key="apply_edits_run", type="primary"):
                    if st.session_state.get("docgen_used_draft_driven"):
                        try:
                            docx_bytes, _ = _run_editor_rebuild_draft_driven(
                                editor_plain,
                                st.session_state["docgen_sample1_bytes"],
                            )
                            st.session_state["docgen_edited_docx_bytes"] = docx_bytes
                            st.success("Document rebuilt. Use **Download edited .docx** below.")
                        except Exception as e:
                            st.error(f"Rebuild failed: {e}")
                    else:
                        stored_blocks = st.session_state.get("docgen_formatted_blocks") or []
                        blocks = _edited_preview_to_blocks(editor_plain, stored_blocks)
                        if not blocks:
                            st.error("No blocks to rebuild. Run formatting first.")
                        else:
                            try:
                                docx_bytes, _ = _run_editor_rebuild(
                                    blocks,
                                    st.session_state["docgen_sample1_bytes"],
                                    overrides if overrides else None,
                                )
                                st.session_state["docgen_edited_docx_bytes"] = docx_bytes
                                st.success("Document rebuilt. Use **Download edited .docx** below.")
                            except Exception as e:
                                st.error(f"Rebuild failed: {e}")
            if st.session_state.get("docgen_edited_docx_bytes"):
                st.download_button(
                    "Download edited .docx",
                    data=st.session_state["docgen_edited_docx_bytes"],
                    file_name="formatted_draft_edited.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_edited_run",
                )
            st.download_button(
                "Download formatted .docx",
                data=st.session_state["docgen_formatted_docx_bytes"],
                file_name="formatted_draft.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                key="dl_formatted_run",
            )
    else:
        st.download_button(
            "Download as .docx",
            data=text_to_docx_bytes(final_draft),
            file_name="draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            key="dl_final_run",
        )


# -----------------------------------------------------------------------------
# Main: run button + persisted final draft (after download/rerun)
# -----------------------------------------------------------------------------

if st.button("Run pipeline (analyze samples → generate draft)", type="primary"):
    run_pipeline()

# Final document draft and download (shown after run and persists after rerun)
# Also show pipeline steps from last run so they don't disappear
if st.session_state.get("docgen_final_draft"):
    st.markdown("---")
    # Pipeline steps recap (from session state) so Step 1–4 remain visible after run
    with st.expander("Pipeline steps (from last run)", expanded=True):
        sections_list_show = st.session_state.get("docgen_sections_list", [])
        if sections_list_show:
            st.subheader("Step 1: Identify sections")
            for i, sec in enumerate(sections_list_show, 1):
                st.markdown(f"**{i}. {sec.get('name', '?')}** — {sec.get('purpose', '') or '(no purpose)'}")
        section_prompts_show = st.session_state.get("docgen_section_prompts", [])
        extracted_show = st.session_state.get("docgen_extracted_samples", [])
        if section_prompts_show:
            st.subheader("Step 2: Prompts and extraction")
            with st.expander("View prompt and required fields per section"):
                for i in range(len(section_prompts_show)):
                    sec_name = sections_list_show[i].get("name", f"Section {i+1}") if i < len(sections_list_show) else f"Section {i+1}"
                    info = section_prompts_show[i] if i < len(section_prompts_show) else {}
                    st.markdown(f"**{sec_name}**")
                    st.caption("Required fields: " + ", ".join(info.get("required_fields", [])))
                    st.text((info.get("prompt", "") or "")[:600] + ("..." if len(info.get("prompt", "")) > 600 else ""))
            if extracted_show:
                with st.expander("View extracted sample text per section"):
                    for i in range(len(extracted_show)):
                        sample_text = extracted_show[i] if i < len(extracted_show) else ""
                        sec_name = sections_list_show[i].get("name", f"Section {i+1}") if i < len(sections_list_show) else f"Section {i+1}"
                        with st.expander(f"{sec_name} — {len(sample_text)} chars"):
                            st.text_area("Extracted sample", value=sample_text[:3000] + ("…" if len(sample_text) > 3000 else ""), height=120, key=f"persisted_extracted_{i}", disabled=True, label_visibility="collapsed")
        field_values_show = st.session_state.get("docgen_field_values", {})
        if field_values_show:
            st.subheader("Step 3: Field values")
            with st.expander("Field → value"):
                for f, v in list(field_values_show.items())[:30]:
                    st.text(f"{f}: {str(v)[:200]}{'…' if len(str(v)) > 200 else ''}")
                if len(field_values_show) > 30:
                    st.caption(f"… and {len(field_values_show) - 30} more fields.")
        draft_parts_show = st.session_state.get("docgen_draft_parts", [])
        if draft_parts_show:
            st.subheader("Step 4: Draft parts")
            st.caption(f"{len(draft_parts_show)} section(s) generated.")
            with st.expander("View draft part per section"):
                for i, part in enumerate(draft_parts_show):
                    sec_name = sections_list_show[i].get("name", f"Section {i+1}") if i < len(sections_list_show) else f"Section {i+1}"
                    with st.expander(f"{sec_name} — {len(part)} chars"):
                        st.text_area("Draft part", value=part[:2500] + ("…" if len(part) > 2500 else ""), height=150, key=f"persisted_draft_part_{i}", label_visibility="collapsed")
        if st.session_state.get("docgen_formatted_preview") is not None:
            st.subheader("Step 5: Format (draft-driven)")
            st.caption("Structure from final draft; styling from sample DOCX.")

    st.markdown("---")
    st.subheader("Final document draft")
    st.text_area(
        "Final draft",
        value=st.session_state["docgen_final_draft"],
        height=420,
        label_visibility="collapsed",
        key="final_draft_show",
    )
    if st.session_state.get("docgen_formatted_preview") is not None:
        st.subheader("Formatted document")
        st.text_area(
            "Formatted preview",
            value=st.session_state["docgen_formatted_preview"],
            height=280,
            label_visibility="collapsed",
            key="formatted_preview_persisted",
        )
        if st.session_state.get("docgen_section_formatting_prompts"):
            with st.expander("Per-section formatting prompts (Step 5a)", expanded=False):
                sections_list_p = st.session_state.get("docgen_sections_list", [])
                for i, sec in enumerate(sections_list_p):
                    prompts_list = st.session_state["docgen_section_formatting_prompts"]
                    instr = prompts_list[i] if i < len(prompts_list) else ""
                    with st.expander(f"{i + 1}. {sec.get('name', 'Section')} — formatting instruction"):
                        st.text_area("Instruction", value=instr, height=180, key=f"section_fmt_{i}_persisted", label_visibility="collapsed")
        if st.session_state.get("docgen_structural_result"):
            sr = st.session_state["docgen_structural_result"]
            with st.expander("Structural pipeline: Rulebook (Step 2)", expanded=False):
                rulebook = sr.get("rulebook") or {}
                layout = rulebook.get("layout_pattern") or []
                st.caption("Sample layout pattern (slot type and style):")
                for i, slot in enumerate(layout[:40]):
                    st.text(f"{i}: {slot.get('section_type')} | {slot.get('block_kind')} | {slot.get('style')}")
                if len(layout) > 40:
                    st.caption(f"… and {len(layout) - 40} more slots.")
            with st.expander("Structural pipeline: Mapping report (Step 3)", expanded=False):
                report = sr.get("mapping_report") or []
                for r in report[:50]:
                    st.text(f"Slot {r.get('slot_index')}: {r.get('template_type')} ← draft {r.get('draft_index')} ({r.get('match_status')}) — {r.get('note', '')}")
                if len(report) > 50:
                    st.caption(f"… and {len(report) - 50} more slots.")
            with st.expander("Structural pipeline: Validation (Step 7)", expanded=False):
                val = sr.get("validation") or {}
                st.write(val.get("message", "—"))
                st.caption(f"Match: {val.get('match')}; same order: {val.get('same_section_order')}; same count: {val.get('same_block_count')}. Sample slots: {val.get('sample_slot_count')}; output slots: {val.get('output_slot_count')}.")
        elif st.session_state.get("docgen_formatting_prompts"):
            with st.expander("How the formatting prompt was generated", expanded=False):
                fp = st.session_state["docgen_formatting_prompts"]
                st.caption("System prompt (sent to the formatting LLM)")
                st.text_area("System", value=fp.get("system", ""), height=220, key="fmt_system_persisted", label_visibility="collapsed")
                st.caption("User message (template structure + style guide + raw text to format)")
                st.text_area("User", value=fp.get("user", ""), height=320, key="fmt_user_persisted", label_visibility="collapsed")
        # Editor: change content or formatting before download
        with st.expander("Edit before download", expanded=True):
            st.caption("Edit the content below and/or set formatting options, then click **Apply edits and download**. For full Word-style editing, download the file and open it in [Word for the Web](https://word.cloud.microsoft) or desktop Word.")
            preview_text_p = st.session_state["docgen_formatted_preview"]
            initial_html_p = _plain_text_to_simple_html(preview_text_p)
            editor_plain_p = preview_text_p
            if _HAS_QUILL:
                st.caption("Use the toolbar for **bold**, *italic*, underline, alignment, lists.")
                quill_content_p = st_quill(
                    value=initial_html_p,
                    html=True,
                    key="docgen_quill_editor_persisted",
                    toolbar=[
                        ["bold", "italic", "underline", "strike"],
                        [{"align": []}],
                        [{"list": "ordered"}, {"list": "bullet"}],
                        ["clean"],
                    ],
                )
                if quill_content_p is not None:
                    st.session_state["docgen_editor_content_html_p"] = quill_content_p
                editor_plain_p = _simple_html_to_plain_text(
                    st.session_state.get("docgen_editor_content_html_p") or initial_html_p
                )
            elif _HAS_LEXICAL:
                st.caption("Edit below (plain text).")
                lexical_content_p = streamlit_lexical(
                    value=preview_text_p,
                    placeholder="Edit document content…",
                    height=280,
                    debounce=500,
                    key="docgen_lexical_editor_persisted",
                )
                if lexical_content_p is not None:
                    st.session_state["docgen_editor_content_plain_p"] = lexical_content_p
                editor_plain_p = st.session_state.get("docgen_editor_content_plain_p") or preview_text_p
            else:
                st.caption("Install **streamlit-quill** for a toolbar editor. Or download and open in [Word for the Web](https://word.cloud.microsoft).")
                st.text_area(
                    "Content (edit as needed)",
                    value=preview_text_p,
                    height=260,
                    key="editor_preview_persisted",
                    label_visibility="collapsed",
                )
                editor_plain_p = st.session_state.get("editor_preview_persisted", preview_text_p)
            col1, col2 = st.columns(2)
            with col1:
                space_after_pt_p = st.number_input(
                    "Space after paragraph (pt); 0 = keep template",
                    min_value=0,
                    max_value=72,
                    value=0,
                    step=1,
                    key="editor_space_after_persisted",
                )
            with col2:
                body_font_size_pt_p = st.number_input(
                    "Body font size (pt); 0 = keep template",
                    min_value=0,
                    max_value=24,
                    value=0,
                    step=1,
                    key="editor_font_size_persisted",
                )
            overrides_p = {}
            if space_after_pt_p > 0:
                overrides_p["space_after_pt"] = space_after_pt_p
            if body_font_size_pt_p > 0:
                overrides_p["body_font_size_pt"] = body_font_size_pt_p
            if st.button("Apply edits and download", key="apply_edits_persisted", type="primary"):
                if st.session_state.get("docgen_used_draft_driven"):
                    try:
                        docx_bytes, _ = _run_editor_rebuild_draft_driven(
                            editor_plain_p,
                            st.session_state["docgen_sample1_bytes"],
                        )
                        st.session_state["docgen_edited_docx_bytes"] = docx_bytes
                        st.success("Document rebuilt. Use **Download edited .docx** below.")
                    except Exception as e:
                        st.error(f"Rebuild failed: {e}")
                else:
                    stored_blocks = st.session_state.get("docgen_formatted_blocks") or []
                    blocks = _edited_preview_to_blocks(editor_plain_p, stored_blocks)
                    if not blocks:
                        st.error("No blocks to rebuild. Run formatting first.")
                    else:
                        try:
                            docx_bytes, _ = _run_editor_rebuild(
                                blocks,
                                st.session_state["docgen_sample1_bytes"],
                                overrides_p if overrides_p else None,
                            )
                            st.session_state["docgen_edited_docx_bytes"] = docx_bytes
                            st.success("Document rebuilt. Use **Download edited .docx** below.")
                        except Exception as e:
                            st.error(f"Rebuild failed: {e}")
        if st.session_state.get("docgen_edited_docx_bytes"):
            st.download_button(
                "Download edited .docx",
                data=st.session_state["docgen_edited_docx_bytes"],
                file_name="formatted_draft_edited.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_edited_persisted",
            )
        st.download_button(
            "Download formatted .docx",
            data=st.session_state["docgen_formatted_docx_bytes"],
            file_name="formatted_draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            key="dl_formatted_persisted",
        )
    else:
        st.download_button(
            "Download as .docx",
            data=text_to_docx_bytes(st.session_state["docgen_final_draft"]),
            file_name="draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            key="dl_final_persisted",
        )