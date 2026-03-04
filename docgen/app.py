import json
import os
import sys
import time
import tempfile
import subprocess
import shutil
from pathlib import Path
from io import BytesIO
import streamlit as st

# -------------------------
# Path setup
# -------------------------
_root = Path(__file__).resolve().parent.parent
_docgen = Path(__file__).resolve().parent
for p in (str(_root), str(_docgen)):
    if p not in sys.path:
        sys.path.insert(0, p)

# -------------------------
# App config
# -------------------------
st.set_page_config(page_title="Document Generator", layout="wide")
st.title("AI Document Generator")

# -------------------------
# Helpers (UNCHANGED)
# -------------------------
def file_to_text(data: bytes, filename: str) -> str:
    if filename and filename.lower().endswith(".docx"):
        from docx import Document
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    return data.decode("utf-8", errors="ignore")


def text_to_docx_bytes(text: str) -> bytes:
    """Build a .docx in memory from plain text (paragraphs split on double newline)."""
    from docx import Document
    doc = Document()
    for block in (text or "").split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block.replace("\n", " "))
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# -------------------------
# Backend (OOP classes)
# -------------------------
from docgen.agents.structure_identification.sectioner import Sectioner
from docgen.agents.content_extraction.extractor import Extractor
from docgen.agents.content_extraction.section_prompt_generator import SectionPromptGenerator
from docgen.agents.data_retrieval.field_fetcher import FieldFetcher, _default_question_for_field
from docgen.agents.data_retrieval.question_generator import QuestionGenerator
from docgen.agents.drafting.section_generator import SectionGenerator
from docgen.agents.drafting.assembler import Assembler
from docgen.agents.review.reviewer import DocumentReviewer
from docgen.agents.category_identification.category_identifier import CategoryIdentifier
from docgen.core.llm_client import LLMClient
from docgen.core.blueprint_manager import BlueprintManager

# -------------------------
# Sidebar
# -------------------------
with st.sidebar:
    # Placeholder for logs at the very top of the sidebar
    logs_placeholder = st.empty()
    
    st.header("Configuration")
    llm_provider = st.radio(
        "LLM Provider", 
        ["Azure", "Gemini"], 
        index=0 if os.getenv("LLM_PROVIDER", "azure").lower() == "azure" else 1
    )
    os.environ["LLM_PROVIDER"] = llm_provider.lower()
    
    if llm_provider == "Gemini":
        st.caption("Using Gemini 3.1 Pro (Preview) with Thinking")
        
    st.divider()

    st.header("Inputs")
    # Blueprint input
    blueprint_name = st.text_input("Blueprint Name", help="Enter a name to save analysis or load an existing blueprint.")
    
    sample1 = st.file_uploader("Sample document 1", type=["txt", "docx"])
    sample2 = st.file_uploader("Sample document 2", type=["txt", "docx"])
    
    curl_input = st.text_area("Case ID or CURL command", height=120)
    firm_id_input = st.number_input("Firm ID (for Case Search)", value=1680, step=1)
    extra_context = st.text_area("Extra context (optional)", height=100)

# -------------------------
# Fade-in CSS (SAFE)
# -------------------------
st.markdown(
    """
    <style>
    .fade-in {
        animation: fadeIn 0.8s ease-in forwards;
        opacity: 0;
        margin-bottom: 12px;
    }
    @keyframes fadeIn {
        to { opacity: 1; }
    }
    .muted-box {
        background-color: #f9fafb;
        padding: 12px;
        border-radius: 6px;
        color: #374151;
        font-size: 0.9rem;
        user-select: none;
    }
    .progress-panel {
        display: flex;
        flex-direction: column;
        padding: 12px 0;
    }
    .progress-panel .fade-in { margin-bottom: 8px; }
    
    /* Increase sidebar width */
    [data-testid="stSidebar"] {
        min-width: 450px !important;
        max-width: 600px !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# -------------------------
# Pipeline
# -------------------------
def run_pipeline(log_status=None, log_container=None):
    _internal_logs = []
    
    def log(title, detail=None):
        if log_status:
            if title:
                log_status.update(label=title)
            if detail:
                if log_container:
                    _internal_logs.append(detail)
                    # Use a CSS trick to force auto-scroll to bottom:
                    # 'column-reverse' anchors to the bottom. We pass the reversed list so DOM is flipped but visually correct.
                    html_lines = [f"<div style='margin-bottom: 8px;'>{l}</div>" for l in reversed(_internal_logs)]
                    html = f"""
                    <div style='max-height: 400px; overflow-y: auto; display: flex; flex-direction: column-reverse; color: #888888; font-size: 0.9em; padding-right: 8px;'>
                        {"".join(html_lines)}
                    </div>
                    """
                    log_container.markdown(html, unsafe_allow_html=True)
                else:
                    log_status.write(detail)

    log("Initializing Pipeline", "Loading LLM Client...")
    llm_client = LLMClient(provider=os.getenv("LLM_PROVIDER", "azure"))
    question_generator = QuestionGenerator(llm_client=llm_client)
    
    # --- BLUEPRINT MANAGER ---
    bpm = BlueprintManager(blueprint_name, _root) if blueprint_name else None
    log("Reading Documents", f"Blueprint: {blueprint_name or 'None'}")
    
    # --- LOAD / READ SAMPLES ---
    sample1_bytes = None
    sample2_bytes = None
    s1_name = "sample1.txt"
    s2_name = "sample2.txt"

    # 1. Try uploads
    if sample1 and sample2:
        try:
            sample1_bytes = sample1.read()
            sample2_bytes = sample2.read()
            s1_name = sample1.name
            s2_name = sample2.name
            
            # Save to blueprint if active
            if bpm:
                bpm.save_bytes(s1_name, sample1_bytes)
                bpm.save_bytes(s2_name, sample2_bytes)
                bpm.save_json("samples_meta.json", {"s1": s1_name, "s2": s2_name})
                st.caption(f"Saved samples to blueprint '{bpm.name}'")
        except Exception as e:
            st.error(f"Could not read uploaded files: {e}")
            return

    # 2. If no uploads, try blueprint
    elif bpm and bpm.has_file("samples_meta.json"):
        meta = bpm.load_json("samples_meta.json")
        if meta and bpm.has_file(meta.get("s1")) and bpm.has_file(meta.get("s2")):
            s1_name = meta["s1"]
            s2_name = meta["s2"]
            sample1_bytes = bpm.load_bytes(s1_name)
            sample2_bytes = bpm.load_bytes(s2_name)
            st.info(f"Loaded samples from Blueprint '{bpm.name}': {s1_name}, {s2_name}")
        else:
            st.warning(f"Blueprint '{bpm.name}' found but samples are missing.")

    if not sample1_bytes or not sample2_bytes:
        st.error("Upload both sample documents (or provide a valid Blueprint with saved samples).")
        return
    
    s1 = file_to_text(sample1_bytes, s1_name)
    s2 = file_to_text(sample2_bytes, s2_name)
    
    # Debug save locally (legacy behavior, can be removed if strictly using blueprints)
    with open("sample1.txt", "wb") as f:
        f.write(str(s1).encode("utf-8"))

    if not (s1 or s2):
        st.error("Both documents are empty. Upload non-empty .txt or .docx files.")
        return

    ctx = (extra_context or "").strip()

    # OOP: docgen pipeline components
    sectioner = Sectioner(llm_client=llm_client)
    extractor = Extractor(llm_client=llm_client)
    section_prompt_generator = SectionPromptGenerator(llm_client=llm_client)
    field_fetcher = FieldFetcher(llm_client=llm_client)
    
    section_generator = SectionGenerator(llm_client=llm_client)
    assembler = Assembler()
    reviewer = DocumentReviewer(llm_client=llm_client)

    # =====================================================
    # STEP 0 — Category Identification
    # =====================================================
    log("Identifying Category...", "Analyzing document types from samples.")
    category_of_document = None
    if bpm:
        category_of_document = bpm.load_text("category.txt")
    
    if not category_of_document:
        category_of_document = CategoryIdentifier(llm_client=llm_client).identify_category(s1, s2)
        if bpm:
            bpm.save_text("category.txt", category_of_document)
    else:
        st.caption("Loaded Category from Blueprint.")

    log("Category Identified", f"Category: {category_of_document}")

    # =====================================================
    # STEP 1 — Section identification (fade-in, slow)
    # =====================================================
    st.subheader("Step 1 · Identifying document sections")
    log("Identifying Sections...", "Dividing samples into logical sections.")

    sections = None
    if bpm:
        sections = bpm.load_json("sections.json")

    if not sections:
        blueprint_data = sectioner.divide_into_sections(s1, s2, category_of_document)
        sections = blueprint_data["sections"]
        if bpm:
            bpm.save_json("sections.json", sections)
    else:
        st.success("Loaded Sections from Blueprint.")
    
    blueprint = {"sections": sections} # ensure blueprint dict structure

    sec_container = st.container()

    for sec in sections:
        with sec_container:
            st.markdown(
                f"""
                <div class="fade-in">
                    <strong>{sec['name']}</strong><br/>
                    <span style="color:#6b7280">{sec.get('purpose','')}</span>
                </div>
                """,
                unsafe_allow_html=True
            )
        time.sleep(0.1) # Faster if loaded

    st.success(f"{len(sections)} sections identified.")

    # =====================================================
    # STEP 2 — Extraction + prompts (continuous loading steps, then arbitrators)
    # =====================================================
    st.subheader("Step 2 · Analyzing sample documents")
    log("Extracting and Generating Prompts", "Extracting sections and generating templates.")

    extracted = None
    prompts = None
    
    if bpm:
        extracted = bpm.load_json("extracted.json")
        prompts = bpm.load_json("prompts.json")

    # If missing or mismatch, re-run
    if not extracted or not prompts or len(extracted) != len(sections) or len(prompts) != len(sections):
        if bpm and (extracted or prompts):
            st.warning("Blueprint analysis incomplete or mismatched. Re-running Step 2...")
            
        step2_status = st.empty()
        step2_status.markdown("**Extracting content from both documents…**")
        step2_details = st.empty()
        def write_step2_details(details: str):
            step2_details.markdown(details)
        extracted = extractor.extract_sections_from_docs(s1, s2, sections, category_of_document, write_step2_details)

        prompts = []
        for i, sec in enumerate(sections):
            sec_name = sec["name"]
            step2_status.markdown(f"**{sec_name}** — extracted text")
            time.sleep(0.1)

            p = section_prompt_generator.generate_prompt_and_fields(
                sec_name,
                sec.get("purpose", ""),
                extracted[i] if i < len(extracted) else "",
                category_of_document
            )
            prompts.append(p)
            step2_status.markdown(f"**{sec_name}** — prompt generated")
            time.sleep(0.1)
        
        step2_status.markdown("**All sections done.**")
        time.sleep(0.4)
        step2_status.empty()
        
        if bpm:
            bpm.save_json("extracted.json", extracted)
            bpm.save_json("prompts.json", prompts)
    else:
        st.success("Loaded Analysis (Extraction & Prompts) from Blueprint.")


    st.success("Extraction and prompts ready.")

    # Arbitrators: one for all extracted text, one for all prompts (default closed)
    with st.expander("**Extracted text (all sections)**", expanded=False):
        for i, sec in enumerate(sections):
            st.markdown(f"### {sec['name']}")
            st.markdown(
                f"<div class='muted-box'>{extracted[i]}</div>",
                unsafe_allow_html=True,
            )
            st.markdown("")
    with st.expander("**Prompts (all sections)**", expanded=False):
        for i, sec in enumerate(sections):
            st.markdown(f"### {sec['name']}")
            st.markdown(
                f"<div class='muted-box'>{prompts[i]['prompt']}</div>",
                unsafe_allow_html=True,
            )
            st.markdown("")

    # =====================================================
    # STEP 3 — Fetch field values via API
    # =====================================================
    all_required = []
    seen = set()
    for i in range(len(sections)):
        for f in prompts[i].get("required_fields", []):
            if f and f not in seen:
                seen.add(f)
                all_required.append(f)
    
    with st.expander("**Required fields**", expanded=False):
        for f in all_required:
            st.markdown(f"**{f}**")

    field_values = {}
    
    # Check if input is a digit (Case ID) or CURL command
    curl_input_clean = (curl_input or "").strip()
    is_case_id = curl_input_clean.isdigit()
    
    if is_case_id:
        # Case 1: Case ID Search
        st.subheader("Step 3: Fetching data from Case Search")
        log("Fetching Data from Search", f"Searching documents for Case ID {curl_input_clean}...")
        if all_required:
            status_placeholder = st.empty()
            progress = st.progress(0, text="Searching documents...")
            
            def on_doc_progress(doc_name: str, index: int, total: int):
                status_placeholder.markdown(f"**Scanning document {index}/{total}:** `{doc_name}`")
                progress.progress(index / total, text=f"Scanning: {doc_name} ({index}/{total})")
                log("Scanning Case Documents", f"Scanning {doc_name} ({index}/{total})")
        
        # --- Capture detailed results (confidence) via callback ---
        field_details = {}
        collected_facts_list = []

        st.markdown("### Live Extraction Progress")
        col1, col2 = st.columns(2)
        with col1:
            live_fields_expander = st.expander("Live Fields Found", expanded=True)
            live_fields_placeholder = live_fields_expander.empty()
            live_fields_placeholder.markdown("*Waiting for first field...*")
        with col2:
            live_facts_expander = st.expander("Live Facts Collected", expanded=True)
            live_facts_placeholder = live_facts_expander.empty()
            live_facts_placeholder.markdown("*Waiting for facts...*")

        def update_live_fields():
            lines = []
            for f, details in field_details.items():
                val_str = str(details['value'])
                if len(val_str) > 150: val_str = val_str[:150] + "..."
                conf = str(details.get('confidence', '')).upper()
                if "HIGH" in conf: conf_badge = "✅"
                elif "PRUNED" in conf: conf_badge = "🚫"
                elif "LOW" in conf: conf_badge = "⚠️"
                else: conf_badge = "❓"
                
                # Using HTML for list item so it formats nicely in the scroll box
                lines.append(f"<li>{conf_badge} <strong>{f}</strong>: {val_str}</li>")
            if lines:
                html = f'<div style="max-height: 300px; overflow-y: auto;"><ul>{"".join(lines)}</ul></div>'
                live_fields_placeholder.markdown(html, unsafe_allow_html=True)

        def update_live_facts():
            if collected_facts_list:
                lines = [f"<li>{fact}</li>" for fact in collected_facts_list]
                html = f'<div style="max-height: 300px; overflow-y: auto;"><ol>{"".join(lines)}</ol></div>'
                live_facts_placeholder.markdown(html, unsafe_allow_html=True)

        def on_field_found_callback(field, value, confidence):
            field_details[field] = {"value": value, "confidence": confidence}
            update_live_fields()

        def on_fact_found_callback(fact):
            collected_facts_list.append(fact)
            update_live_facts()

        # Fetch using the new search strategy
        field_values = field_fetcher.fetch_fields_from_case_search(
            case_id=curl_input_clean,
            firm_id=firm_id_input,
            required_fields=all_required,
            on_doc_start=on_doc_progress,
            on_field_found=on_field_found_callback,
            on_fact_found=on_fact_found_callback,
            category_of_document=category_of_document
        )
        
        status_placeholder.markdown(f"**Done.** Found {len(field_values)} of {len(all_required)} fields.")
        progress.progress(1.0, text="Done.")

        # --- Group fields by confidence ---
        high_conf = []
        low_conf = []
        pruned_conf = []
        unrecognized = []
        
        for f in all_required:
            if f in field_details:
                conf = str(field_details[f]["confidence"]).upper()
                val = field_details[f]["value"]
                if "HIGH" in conf:
                    high_conf.append((f, val))
                elif "PRUNED" in conf:
                    pruned_conf.append((f, val))
                else:
                    low_conf.append((f, val))
            elif f not in field_values:
                 # If it wasn't captured in callback AND not in final result dict
                 unrecognized.append(f)
            else:
                 # In result dict but missed callback? Treat as low confidence fallback
                 low_conf.append((f, field_values[f]))

        # --- Display in Expanders ---
        st.markdown("### Extraction Results")
        
        with st.expander(f"High Confidence ({len(high_conf)})", expanded=True):
            if high_conf:
                for f, v in high_conf:
                    vstr = str(v)
                    st.markdown(f"**{f}**: {vstr[:100]}...")
            else:
                st.caption("No high confidence matches.")

        with st.expander(f"Irrelevant / Pruned ({len(pruned_conf)})", expanded=True):
            if pruned_conf:
                for f, v in pruned_conf:
                    vstr = str(v)
                    st.markdown(f"🚫 **{f}**: {vstr}")
            else:
                st.caption("No fields pruned.")

        with st.expander(f"Low Confidence / Ambiguous ({len(low_conf)})", expanded=False):
            if low_conf:
                for f, v in low_conf:
                    vstr = str(v)
                    st.markdown(f"**{f}**: {vstr[:100]}...")
            else:
                st.caption("No low confidence matches.")

        with st.expander(f"Unrecognized Fields ({len(unrecognized)})", expanded=False):
            if unrecognized:
                for f in unrecognized:
                    st.markdown(f"- {f}")
            else:
                st.caption("All fields found!")


    elif curl_input_clean:
        # Case 2: CURL Command (Legacy Chat API)
        st.subheader("Step 3: Fetching data via API")
        if all_required:
            
            # Try load questions from blueprint
            field_to_question = None
            if bpm:
                field_to_question = bpm.load_json("generated_questions.json")

            if not field_to_question:
                with st.status("Generating questions for each field...", state="running"):
                    field_to_question = question_generator.generate_questions_for_fields(all_required,category_of_document)
                if bpm:
                    bpm.save_json("generated_questions.json", field_to_question)
            else:
                st.caption("Loaded Questions from Blueprint.")

            generated_questions_file = "generated_questions.json"
            with open(generated_questions_file, "wb") as f:
                f.write(json.dumps(field_to_question).encode("utf-8"))
            st.success(f"Generated {len(field_to_question)} questions.")
            with st.expander("Field → question used for API", expanded=False):
                for f, q in field_to_question.items():
                    st.markdown(f"**{f}** → \"{q}\"")

            first_field = all_required[0]
            first_question = field_to_question.get(first_field) or _default_question_for_field(first_field)
            debug = field_fetcher.call_chat_api_with_question_debug((curl_input or "").strip(), first_question)
            with st.expander("API test (first request)", expanded=False):
                st.caption(f"Question: \"{first_question}\"")
                if debug.get("error"):
                    st.error(debug["error"])
                st.code(f"Status: {debug.get('status_code')} | Response keys: {debug.get('response_keys', [])}")
                if debug.get("extracted_preview"):
                    st.success(f"Extracted: {debug['extracted_preview']}")
                else:
                    st.warning("No text could be extracted from the response. Check that your API returns a body with one of: content, answer, message, text, or choices[0].message.content.")

            status_placeholder = st.empty()
            progress = st.progress(0, text="Preparing...")

            def on_field_start(field_name: str, index: int, total: int):
                status_placeholder.markdown(f"**Fetching field {index} of {total}:** `{field_name}`")
                progress.progress(index / total, text=f"Fetching: {field_name} ({index}/{total})")

            field_values = field_fetcher.fetch_all_fields_via_chat(
                (curl_input or "").strip(), all_required, field_to_question, on_field_start=on_field_start
            )
            status_placeholder.markdown("**Done.** All fields fetched.")
            progress.progress(1.0, text="Done.")

            status_placeholder.markdown("**Fetching case summary…**")
            case_summary = field_fetcher.fetch_case_summary((curl_input or "").strip())
            field_values["case_summary"] = case_summary or ""
            if case_summary:
                st.caption("Case summary received and will be passed to each section.")
            status_placeholder.markdown("**Done.** All fields fetched.")

        with st.expander("Field → value (from API)", expanded=False):
            for f, v in field_values.items():
                vstr = str(v)
                st.markdown(f"**{f}** → {vstr[:200] + '...' if len(vstr) > 200 else vstr or '(empty)'}")
    else:
        st.subheader("Step 3: Field data")
        st.info("No CURL provided. Using extra context only if provided.")

    if ctx:
        field_values["case_summary_or_context"] = ctx
        field_values["extra_context"] = ctx

    # =====================================================
    # STEP 4 — Drafting (text_area, append-only, no headings)
    # =====================================================
    st.subheader("Step 4 · Drafting the document")
    log("Drafting Document", "Generating document sections...")

    left, right = st.columns([3, 1])

    draft_text = ""
    draft_sections = []  # for Step 5 assemble()
    draft_height = 220  # will grow as content grows
    total = len(sections)
    completed_sections = []

    with left:
        draft_box = st.empty()  # single box, content updated each step
    with right:
        right_placeholder = st.empty()  # progress panel, height matches draft box

    for i, sec in enumerate(sections, start=1):
        section_name = sec["name"]
        log(f"Drafting Section {i}/{total}", f"Working on: {section_name}")
        
        req_fields = prompts[i - 1].get("required_fields", [])
        section_field_values = {f: field_values.get(f, "") for f in field_values} # passing all the things in case
        if ctx:
            section_field_values["case_summary_or_context"] = ctx
        if field_values.get("case_summary"):
            section_field_values["case_summary"] = field_values["case_summary"]

        section_text = section_generator.generate_section(
            prompts[i - 1]["prompt"],
            section_field_values,
            sample_text=extracted[i - 1] if i - 1 < len(extracted) else "",
            section_name=section_name,
        )

        # Check for Omission
        is_skipped = "<<SECTION_SKIPPED>>" in section_text
        if is_skipped:
            section_text = "" # Clear text for assembly
            display_name = f"🚫 {section_name} (Skipped)"
        else:
            draft_text += section_text.strip() + "\n\n"
            display_name = section_name

        draft_sections.append(section_text)

        # ---- Increase height gradually
        draft_height = min(draft_height + 120, 900)

        # ---- Right panel: same height as draft box, progress (fade-in), completed (bold)
        # Inline animation so it runs when placeholder content is replaced
        fade_style = "opacity:0; animation: fadeIn 0.8s ease-in forwards; margin-bottom: 8px;"
        completed_html = "".join(f"<div>✓ {c}</div>" for c in completed_sections)
        right_placeholder.markdown(
            f"""
            <div class="progress-panel" style="min-height:{draft_height}px;">
                <div class="fade-in" style="{fade_style}"><strong>Drafting section {i} of {total}</strong></div>
                <div class="fade-in" style="{fade_style}"><strong>{display_name}</strong></div>
                {completed_html}
            </div>
            """,
            unsafe_allow_html=True,
        )
        completed_sections.append(display_name)

        # ---- Left panel: single box
        with left:
            draft_box.text_area(
                label="Draft being generated",
                value=draft_text,
                height=draft_height,
                key=f"draft_live_{i}",
                disabled=True,
            )

        time.sleep(0.4)

    # ---- Final right panel: all sections with ✓ (including last)
    completed_html = "".join(f"<div><strong>✓ {c}</strong></div>" for c in completed_sections)
    right_placeholder.markdown(
        f"""
        <div class="progress-panel" style="min-height:{draft_height}px;">
            <div style="margin-bottom:8px;">Drafting complete.</div>
            {completed_html}
        </div>
        """,
        unsafe_allow_html=True,
    )

    # =====================================================
    # STEP 5 — Review, Polishing and formatting
    # =====================================================
    st.subheader("Step 5 · Reviewing, Polishing and formatting the document")
    log("Reviewing and Polishing", "Assembling document and reviewing flow...")

    # 1. Assemble
    initial_draft = assembler.assemble(blueprint, draft_sections)
    
    # 2. Review
    final_draft = initial_draft
    with st.spinner("Reviewing draft for consistency and flow..."):
        try:
            reviewed_draft = reviewer.review_draft(initial_draft, field_values, category_of_document)
            if reviewed_draft and len(reviewed_draft) > 100:
                final_draft = reviewed_draft
                st.success("Draft reviewed and polished.")
                log("Review Complete", "Draft successfully polished.")
            else:
                st.warning("Reviewer returned empty or too short text. Using initial draft.")
                log("Review Skipped", "Using initial draft due to reviewer error.")
        except Exception as e:
            st.error(f"Review step failed: {e}. Using initial draft.")
            log("Review Failed", f"Error: {e}")

    formatted_docx_bytes = None
    formatting_error = None

    # Use sample 1 as formatting template when it is a DOCX (run in subprocess so formatting's "utils" package is used)
    if s1_name and s1_name.lower().endswith(".docx") and sample1_bytes:
        log("Formatting Document", "Applying DOCX template styles...")
        with st.status("Applying template formatting (styles & structure from Sample 1)…", state="running"):
            _formatting_dir = _root / "formatting"
            try:
                with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as _draft_f:
                    _draft_f.write(final_draft)
                    _draft_path = _draft_f.name
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as _tpl_f:
                    _tpl_f.write(sample1_bytes)
                    _tpl_path = _tpl_f.name
                _env = os.environ.copy()
                _env["PYTHONPATH"] = str(_formatting_dir)
                _code = """
import os, sys
sys.path.insert(0, os.environ['FORMATTING_DIR'])
from backend import process_document
draft_path = os.environ['DRAFT_PATH']
tpl_path = os.environ['TEMPLATE_PATH']
with open(draft_path, 'r', encoding='utf-8') as f:
    draft = f.read()
with open(tpl_path, 'rb') as f:
    out_path, _ = process_document(draft, f)
print(out_path)
"""
                _env["FORMATTING_DIR"] = str(_formatting_dir)
                _env["DRAFT_PATH"] = _draft_path
                _env["TEMPLATE_PATH"] = _tpl_path
                _r = subprocess.run(
                    [sys.executable, "-c", _code],
                    cwd=str(_formatting_dir),
                    env=_env,
                    capture_output=True,
                    text=True,
                    timeout=300,
                )
                for _p in (_draft_path, _tpl_path):
                    try:
                        os.unlink(_p)
                    except OSError:
                        pass
                if _r.returncode == 0 and _r.stdout:
                    _out_path = _r.stdout.strip()
                    if os.path.isfile(_out_path):
                        with open(_out_path, "rb") as _f:
                            formatted_docx_bytes = _f.read()
                else:
                    formatting_error = _r.stderr or _r.stdout or "Formatting subprocess failed"
            except Exception as _e:
                formatting_error = str(_e)
        if formatting_error:
            st.warning(f"Formatting pipeline failed (using plain draft): {formatting_error}")

    # --- Save output to generated_documents folder ---
    try:
        # Determine folder structure: generated_documents/{firm_id}/{case_id}/
        
        # Determine Case ID
        save_case_id = "unknown_case"
        if is_case_id:
            save_case_id = str(curl_input_clean)
        elif curl_input_clean:
            # Try to extract c_matter_id from curl payload if possible
            # Simple regex search for c_matter_id in the curl string
            import re
            m = re.search(r'c_matter_id["\']?\s*[:=]\s*["\']?(\d+)', curl_input_clean)
            if m:
                save_case_id = m.group(1)
            else:
                # Fallback: simple hash or timestamp if no ID found
                save_case_id = f"custom_curl_{int(time.time())}"
        
        save_firm_id = str(firm_id_input) if firm_id_input else "default_firm"
        
        output_dir = _root / "generated_documents" / save_firm_id / save_case_id
        output_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        doc_name = f"generated_draft_{timestamp}.docx"
        output_path = output_dir / doc_name
        
        # Save the bytes
        final_bytes = formatted_docx_bytes if formatted_docx_bytes else text_to_docx_bytes(final_draft)
        
        with open(output_path, "wb") as f:
            f.write(final_bytes)
            
        st.success(f"Document saved to: `{output_path}`")
        
    except Exception as e:
        st.error(f"Failed to save document to folder: {e}")

    st.subheader("Formatted document")

    st.text_area(
        "Final output",
        value=final_draft,
        height=600,
        label_visibility="collapsed"
    )

    if formatted_docx_bytes:
        st.download_button(
            "Download formatted .docx",
            data=formatted_docx_bytes,
            file_name="formatted_draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    else:
        st.download_button(
            "Download .docx",
            data=text_to_docx_bytes(final_draft),
            file_name="draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )

    # Persist results so they stay visible after download or other re-runs
    st.session_state["pipeline_done"] = True
    st.session_state["pipeline_blueprint"] = blueprint
    st.session_state["pipeline_sections"] = sections
    st.session_state["pipeline_extracted"] = extracted
    st.session_state["pipeline_prompts"] = prompts
    st.session_state["pipeline_draft_text"] = draft_text
    st.session_state["pipeline_final_draft"] = final_draft
    st.session_state["pipeline_completed_sections"] = completed_sections
    st.session_state["pipeline_formatted_docx_bytes"] = formatted_docx_bytes  # None or bytes


def render_saved_pipeline_results():
    """Re-render all steps and final output from session state (keeps UI after download)."""
    sections = st.session_state.get("pipeline_sections", [])
    extracted = st.session_state.get("pipeline_extracted", [])
    prompts = st.session_state.get("pipeline_prompts", [])
    draft_text = st.session_state.get("pipeline_draft_text", "")
    final_draft = st.session_state.get("pipeline_final_draft", "")
    completed_sections = st.session_state.get("pipeline_completed_sections", [])
    formatted_docx_bytes = st.session_state.get("pipeline_formatted_docx_bytes")

    if not sections:
        return

    # Step 1
    st.subheader("Step 1 · Identifying document sections")
    for sec in sections:
        st.markdown(
            f"""
            <div class="fade-in">
                <strong>{sec['name']}</strong><br/>
                <span style="color:#6b7280">{sec.get('purpose','')}</span>
            </div>
            """,
            unsafe_allow_html=True,
        )
    st.success(f"{len(sections)} sections identified.")

    # Step 2
    st.subheader("Step 2 · Analyzing sample documents")
    with st.expander("**Extracted text (all sections)**", expanded=False):
        for i, sec in enumerate(sections):
            st.markdown(f"### {sec['name']}")
            if i < len(extracted):
                st.markdown(
                    f"<div class='muted-box'>{extracted[i]}</div>",
                    unsafe_allow_html=True,
                )
            st.markdown("")
    with st.expander("**Prompts (all sections)**", expanded=False):
        for i, sec in enumerate(sections):
            st.markdown(f"### {sec['name']}")
            if i < len(prompts):
                st.markdown(
                    f"<div class='muted-box'>{prompts[i].get('prompt', '')}</div>",
                    unsafe_allow_html=True,
                )
            st.markdown("")
    st.success("Extraction and prompts ready.")

    # Step 4 (summary)
    st.subheader("Step 4 · Drafting the document")
    left, right = st.columns([3, 1])
    with left:
        st.text_area(
            "Draft being generated",
            value=draft_text,
            height=min(220 + 120 * max(0, len(sections) - 1), 900),
            key="saved_draft_display",
            disabled=True,
        )
    with right:
        completed_html = "".join(f"<div><strong>✓ {c}</strong></div>" for c in completed_sections)
        st.markdown(
            f"""
            <div class="progress-panel">
                <div style="margin-bottom:8px;">Drafting complete.</div>
                {completed_html}
            </div>
            """,
            unsafe_allow_html=True,
        )

    # Step 5
    st.subheader("Step 5 · Polishing and formatting the document")
    st.subheader("Formatted document")
    st.text_area(
        "Final output",
        value=final_draft,
        height=600,
        key="saved_final_display",
        label_visibility="collapsed",
    )
    if formatted_docx_bytes:
        st.download_button(
            "Download formatted .docx",
            data=formatted_docx_bytes,
            file_name="formatted_draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            key="download_formatted_saved",
        )
    else:
        st.download_button(
            "Download .docx",
            data=text_to_docx_bytes(final_draft),
            file_name="draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            key="download_plain_saved",
        )


# -------------------------
# Run
# -------------------------
if st.button("Run document generation", type="primary"):
    st.session_state.pop("pipeline_done", None)  # clear so we run fresh
    status = st.empty()
    
    with logs_placeholder.container():
        st.markdown("### Process Logs")
        log_status = st.status("Initializing...", expanded=True)
        # Use an empty placeholder so we can overwrite it and force scroll-to-bottom
        log_container = log_status.empty()
        
    try:
        run_pipeline(log_status, log_container)
    except Exception as e:
        status.empty()
        st.error("Pipeline failed. See details below.")
        st.exception(e)
        log_status.update(label="Pipeline Failed", state="error")
    else:
        status.empty()
        log_status.update(label="Pipeline Complete", state="complete")
        
elif st.session_state.get("pipeline_done"):
    render_saved_pipeline_results()
    with logs_placeholder.container():
        st.markdown("### Process Logs")
        st.info("Pipeline completed. Logs are only available during active generation.")
