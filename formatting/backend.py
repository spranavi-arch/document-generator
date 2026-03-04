import base64
import json
import logging
import os
import tempfile

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

from utils.docx_to_images import docx_to_page_images, docx_to_page_images_base64, ocr_page_images
from utils.formatter import (
    clear_document_body,
    force_legal_run_format_document,
    force_single_column,
    inject_blocks,
    remove_trailing_empty_and_noise,
)
from utils.llm_formatter import format_text_with_llm
from utils.style_extractor import (
    _paragraph_has_bottom_border,
    extract_document_blueprint,
    extract_styles,
    load_extracted_styles,
    save_document_blueprint,
    save_extracted_styles,
)

# Summons-style page margins (generous, like formal legal documents)
DEFAULT_TOP_MARGIN_IN = 1.0
DEFAULT_BOTTOM_MARGIN_IN = 1.0
DEFAULT_LEFT_MARGIN_IN = 1.25
DEFAULT_RIGHT_MARGIN_IN = 1.25


def _apply_default_margins(doc):
    """Ensure every section has at least default wide margins (proper spacing from page edges)."""
    try:
        for section in doc.sections:
            section.top_margin = Inches(DEFAULT_TOP_MARGIN_IN)
            section.bottom_margin = Inches(DEFAULT_BOTTOM_MARGIN_IN)
            section.left_margin = Inches(DEFAULT_LEFT_MARGIN_IN)
            section.right_margin = Inches(DEFAULT_RIGHT_MARGIN_IN)
    except Exception:
        pass


def _project_dir():
    return os.path.dirname(os.path.abspath(__file__))


def _tighten_footer_spacing(doc):
    """Tighten spacing on the last page footer that starts with SUPERIOR COURT / NEW HAVEN COUNTY.

    We set space_before/space_after to 0 and use single line spacing from that footer
    heading through the end of the document, so the summons-and-verified-complaint
    footer page matches the sample's tight legal layout.
    """
    if not doc or not getattr(doc, "paragraphs", None):
        return
    start_idx = None
    try:
        for i, para in enumerate(doc.paragraphs):
            t = (para.text or "").strip().upper()
            if "SUPERIOR COURT" in t and "NEW HAVEN COUNTY" in t:
                start_idx = i
        if start_idx is None:
            return
        for para in doc.paragraphs[start_idx:]:
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.line_spacing = 1
        # Also tighten spacing inside tables (footer right column, etc.)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        pf = para.paragraph_format
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(0)
                        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        pf.line_spacing = 1
    except Exception:
        # Spacing is a visual enhancement only; never fail the pipeline over it.
        return


def _log_blocks(blocks, project_dir: str):
    """Persist the LLM/segmenter output blocks for debugging (good vs bad runs)."""
    try:
        out_dir = os.path.join(project_dir, "output")
        os.makedirs(out_dir, exist_ok=True)
        path = os.path.join(out_dir, "last_blocks.json")
        data = [
            {"block_type": bt, "text": (text or "")}
            for (bt, text) in (blocks or [])
        ]
        with open(path, "w", encoding="utf-8") as f:
            json.dump({"blocks": data, "num_blocks": len(data)}, f, indent=2, ensure_ascii=False)
    except Exception:
        # Logging must never break formatting; ignore failures.
        return


def _log_verification_presence(blocks: list):
    """Log whether Attorney's Verification (or similar) appears in blocks to help debug missing/blacked-out bottom sections.
    If absent, the LLM may have truncated (e.g. max_tokens) or the raw input may not include that section."""
    try:
        if not blocks:
            logging.info("Verification tracking: no blocks")
            return
        n = len(blocks)
        has_verification_heading = any(
            "attorney" in (t or "").lower() and "verification" in (t or "").lower()
            for _, t in blocks
        )
        has_request_for_claim = any(
            "request for claim" in (t or "").lower() or "notice of entry" in (t or "").lower()
            for _, t in blocks
        )
        logging.info(
            "Verification tracking: blocks=%s, verification_heading=%s, request_for_claim_or_notice=%s",
            n, has_verification_heading, has_request_for_claim,
        )
    except Exception:
        return


def _get_document_font_from_schema(schema: dict) -> str:
    """Pick one font name from the template schema so the whole document uses that font. Prefer paragraph/normal style."""
    style_formatting = schema.get("style_formatting") or {}
    style_map = schema.get("style_map") or {}
    for style_key in (style_map.get("paragraph"), "Normal", "Body Text", "List Paragraph"):
        if not style_key:
            continue
        rf = (style_formatting.get(style_key) or {}).get("run_format") or {}
        if rf.get("name"):
            return str(rf["name"]).strip()
    for fmt in style_formatting.values():
        rf = (fmt or {}).get("run_format") or {}
        if rf.get("name"):
            return str(rf["name"]).strip()
    return "Times New Roman"


def get_document_preview_text(docx_path: str) -> str:
    """Build a plain-text preview of the formatted DOCX for display before download.
    Paragraphs with only a bottom border (section underlines) are emitted as [SECTION_UNDERLINE]."""
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text and _paragraph_has_bottom_border(para):
            lines.append("[SECTION_UNDERLINE]")
        else:
            lines.append(text if text else "")
    return "\n\n".join(lines).strip()


def extract_and_store_styles(template_file) -> dict:
    """Extract styles from the uploaded DOCX and save to JSON. Returns the style schema."""
    doc = Document(template_file)
    schema = extract_styles(doc)
    save_extracted_styles(schema, base_dir=_project_dir())
    blueprint = extract_document_blueprint(doc)
    save_document_blueprint(blueprint, base_dir=_project_dir())
    return schema


def process_document(generated_text, template_file, use_slot_fill: bool = False):
    """
    Input 1: Uploaded DOCX template (desired styles and formatting).
    Input 2: Raw legal text (unformatted).
    When use_slot_fill=True and template has template_structure: LLM fills exact slots; inject_blocks uses template_structure.
    When use_slot_fill=False: segment entire text using template styles; template page images sent to LLM when available.
    """
    project_dir = _project_dir()
    doc = Document(template_file)
    _apply_default_margins(doc)

    schema = extract_styles(doc)
    save_extracted_styles(schema, base_dir=project_dir)

    # Always generate template page images so the LLM can use them when the LLM path is used
    try:
        template_file.seek(0)
        data = template_file.read()
        template_file.seek(0)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(data)
            tmp.flush()
            template_path = tmp.name
        doc_for_images = Document(template_path)
        force_single_column(doc_for_images)
        fd, single_column_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc_for_images.save(single_column_path)
        page_bytes = docx_to_page_images(single_column_path, dpi=150, max_pages=15)
        if page_bytes:
            schema["template_page_images"] = [base64.b64encode(b).decode("ascii") for b in page_bytes]
            logging.info("Template page images: %s pages", len(schema["template_page_images"]))
            template_page_ocr_texts = ocr_page_images(page_bytes)
            if template_page_ocr_texts and any(t.strip() for t in template_page_ocr_texts):
                schema["template_page_ocr_texts"] = template_page_ocr_texts
        else:
            logging.info("Template page images: none (conversion failed or no LibreOffice)")
        for path in (single_column_path, template_path):
            if path and os.path.isfile(path):
                try:
                    os.unlink(path)
                except OSError:
                    pass
    except Exception:
        logging.info("Template page images: none (conversion failed or no LibreOffice)")

    blocks = []
    try:
        num_images = len(schema.get("template_page_images") or [])
        logging.info("Sending to LLM: template_page_images=%s pages", num_images)
        blocks = format_text_with_llm(
            generated_text,
            schema,
            use_slot_fill=use_slot_fill,
            template_page_images=schema.get("template_page_images") or [],
            template_page_ocr_texts=schema.get("template_page_ocr_texts") or [],
        )
    except Exception:
        pass

    # If LLM returned no blocks or only blocks with empty text, inject the raw draft so the DOCX is never blank
    has_any_content = any((t or "").strip() for _, t in blocks) if blocks else False
    if (not blocks or not has_any_content) and (generated_text or "").strip():
        para_style = (schema.get("style_map") or {}).get("paragraph") or "Normal"
        blocks = [(para_style, (generated_text or "").strip())]

    # Write block list to disk so we can diff "good" vs "bad" runs for the same input.
    _log_blocks(blocks, project_dir)

    # Track whether Attorney's Verification / bottom section is present (helps debug missing or blacked-out sections)
    _log_verification_presence(blocks)

    clear_document_body(doc)
    # Preserve template's section/column layout (do not force single-column so two-column claimant/attorney blocks match template)
    template_structure = schema.get("template_structure") if use_slot_fill else None
    inject_blocks(
        doc,
        blocks,
        style_map=schema["style_map"],
        style_formatting=schema.get("style_formatting", {}),
        line_samples=schema.get("line_samples", []),
        section_heading_samples=schema.get("section_heading_samples", []),
        template_structure=template_structure,
        numbered_num_id=schema.get("numbered_num_id"),
        numbered_ilvl=schema.get("numbered_ilvl", 0),
        bold_phrases_from_template=schema.get("bold_phrases_from_template"),
        caption_table_layout=schema.get("caption_table_layout"),
    )
    document_font = _get_document_font_from_schema(schema)
    force_legal_run_format_document(doc, font_name=document_font)
    _tighten_footer_spacing(doc)
    remove_trailing_empty_and_noise(doc)

    output_path = os.path.join(project_dir, "output", "formatted_output.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    preview_text = get_document_preview_text(output_path)
    return output_path, preview_text