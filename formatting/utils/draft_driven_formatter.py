"""
Draft-driven formatting: structure comes from the final draft, styling from the sample.

- Step 1: Parse the final draft into blocks (heading, body, numbered, line, etc.) using the
  same structure as the draft (paragraph breaks = double newline; no template slot count).
- Step 2: From the sample DOCX, use only the style guide: which style to use for each role
  (heading, body, numbered, etc.) and that style's formatting (spacing, font, alignment).
- Step 3: Build the output DOCX by adding one paragraph per draft block, applying the
  sample's style for that block's role. Result: same structure as the draft, same look as sample.
"""

import re
from docx import Document

from utils.formatter import (
    _apply_paragraph_format,
    _apply_run_format,
    _render_checkboxes,
    clear_document_body,
    force_single_column,
    remove_trailing_empty_and_noise,
)

try:
    from utils.html_to_docx import _paragraph_border_bottom
except Exception:
    _paragraph_border_bottom = None


# Block roles that map to style_map keys
ROLE_HEADING = "heading"
ROLE_SECTION_HEADER = "section_header"
ROLE_PARAGRAPH = "paragraph"
ROLE_NUMBERED = "numbered"
ROLE_LINE = "line"
ROLE_SIGNATURE_LINE = "signature_line"
ROLE_SECTION_UNDERLINE = "section_underline"


def parse_draft_into_blocks(draft_text: str) -> list[tuple[str, str]]:
    """
    Parse the final draft into a list of (role, text) blocks.
    Structure is driven only by the draft: double newline = paragraph break.
    No template slot count; no reordering.
    """
    if not draft_text or not draft_text.strip():
        return []
    blocks = []
    for para in draft_text.strip().split("\n\n"):
        para = (para or "").strip()
        if not para:
            continue
        if para == "[SECTION_UNDERLINE]":
            blocks.append((ROLE_SECTION_UNDERLINE, ""))
            continue
        # Signature line: mostly underscores
        if _is_signature_line(para):
            blocks.append((ROLE_SIGNATURE_LINE, para))
            continue
        # Separator line: dashes/dots ending in X
        if _is_separator_line(para):
            blocks.append((ROLE_LINE, para))
            continue
        # Numbered list item: starts with "1. " or "1) " or "a. " etc.
        if _is_numbered_item(para):
            blocks.append((ROLE_NUMBERED, para))
            continue
        # Section heading: short, often ALL CAPS (e.g. "SUMMONS", "AS AND FOR A FIRST CAUSE OF ACTION")
        if _is_section_heading(para):
            blocks.append((ROLE_SECTION_HEADER, para))
            continue
        # Document title / main heading: very short, often centered in samples
        if _is_document_title(para):
            blocks.append((ROLE_HEADING, para))
            continue
        # Default: body paragraph
        blocks.append((ROLE_PARAGRAPH, para))
    return blocks


def _is_signature_line(text: str) -> bool:
    if not text or len(text) > 80:
        return False
    s = text.strip()
    if "____" in s or (len(s) < 30 and s.count("_") >= 3):
        return True
    return False


def _is_separator_line(text: str) -> bool:
    s = text.strip()
    if len(s) < 10:
        return False
    if not s.endswith("X") and not s.endswith("x"):
        return False
    rest = s[:-1].strip()
    return all(c in " \t\-_.=\u00A0" for c in rest)


def _is_numbered_item(text: str) -> bool:
    s = text.strip()
    # "1. ", "1) ", "a. ", "i. ", "(1) "
    if re.match(r"^\d+[.)]\s+", s):
        return True
    if re.match(r"^[a-z][.)]\s+", s) and len(s) > 4:
        return True
    if re.match(r"^[ivxIVX]+[.)]\s+", s) and len(s) > 4:
        return True
    if re.match(r"^\(\d+\)\s+", s):
        return True
    return False


def _is_section_heading(text: str) -> bool:
    """Short line that looks like a section title (e.g. ALL CAPS cause of action)."""
    s = text.strip()
    if len(s) > 100:
        return False
    # ALL CAPS and looks like a title
    if len(s) > 2 and sum(1 for c in s if c.isalpha() and c.isupper()) / max(1, sum(1 for c in s if c.isalpha())) >= 0.8:
        return True
    # Known phrases
    lower = s.lower()
    if "as and for a" in lower and "cause of action" in lower:
        return True
    if "jury trial demanded" in lower:
        return True
    return False


def _is_document_title(text: str) -> bool:
    """Very short line, often document type (SUMMONS, COMPLAINT)."""
    s = text.strip()
    if len(s) > 60:
        return False
    lower = s.lower()
    if lower in ("summons", "complaint", "verified complaint", "notice of motion", "affidavit"):
        return True
    if len(s) <= 30 and s.isupper():
        return True
    return False


def build_document_from_draft(
    doc: Document,
    blocks: list[tuple[str, str]],
    style_map: dict,
    style_formatting: dict,
) -> None:
    """
    Add one paragraph per block to the document, using the sample's style for each block's role.
    doc: template document (body will be cleared first).
    blocks: list of (role, text) from parse_draft_into_blocks().
    style_map: role -> style name (e.g. {"heading": "Heading 1", "paragraph": "Normal"}).
    style_formatting: style name -> {paragraph_format, run_format}.
    """
    clear_document_body(doc)
    force_single_column(doc)
    for role, text in blocks:
        style_name = style_map.get(role) or style_map.get("paragraph") or "Normal"
        try:
            _ = doc.styles[style_name]
        except (KeyError, AttributeError):
            style_name = "Normal"
        fmt = style_formatting.get(style_name) or {}
        pf = fmt.get("paragraph_format") or {}
        rf = fmt.get("run_format") or {}
        text = (text or "").strip()
        if role == ROLE_SECTION_UNDERLINE:
            p = doc.add_paragraph(style=style_name)
            if _paragraph_border_bottom:
                _paragraph_border_bottom(p, pt=0.5)
            _apply_paragraph_format(p, pf)
            continue
        if role == ROLE_LINE:
            line_text = text or "----------------------------------------------------------------------X"
            p = doc.add_paragraph(style=style_name)
            run = p.add_run(line_text)
            _apply_run_format(run, rf)
            _apply_paragraph_format(p, pf)
            continue
        if role == ROLE_SIGNATURE_LINE:
            line_text = text or "_________________________"
            p = doc.add_paragraph(style=style_name)
            run = p.add_run(line_text)
            _apply_run_format(run, rf)
            _apply_paragraph_format(p, pf)
            continue
        # heading, section_header, paragraph, numbered
        p = doc.add_paragraph(style=style_name)
        if text:
            run = p.add_run(_render_checkboxes(text))
            _apply_run_format(run, rf)
        _apply_paragraph_format(p, pf)
    remove_trailing_empty_and_noise(doc)
