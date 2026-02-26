"""
Phase 1: Convert sample DOCX into template by replacing dynamic text with placeholders.

Run-safe: replace at PARAGRAPH level. Full paragraph text = concat all runs; apply replacements;
rebuild paragraph into a SINGLE run. Do not change paragraph style/alignment/tabstops/borders.

Structural rules (not exact match only):
- Caption: ALL CAPS name before "Plaintiff" -> {{PLAINTIFF_NAME}}; before "Defendant" -> {{DEFENDANT_NAME}}
- Replace ALL occurrences of those names (caption + body, e.g. "Plaintiff, <NAME>")
- Index No.: value -> {{INDEX_NO}}; Date Filed: value -> {{DATE_FILED}}
- Apply to doc.paragraphs and all table cell paragraphs. Replace ALL occurrences.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from utils.placeholder_docx import paragraph_full_text, rebuild_paragraph_single_run

log = logging.getLogger(__name__)

# --- Iterate all paragraphs (body + table cells) ---
def _iter_body_blocks(doc: Document):
    """Yield (paragraph, table_id, row, col) in document order."""
    body = doc.element.body
    qp, qt = qn("w:p"), qn("w:tbl")
    table_id = 0
    for child in body.iterchildren():
        if child.tag == qp:
            yield Paragraph(child, doc), None, None, None
        elif child.tag == qt:
            tbl = Table(child, doc)
            for ri, row in enumerate(tbl.rows):
                for ci, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        yield para, table_id, ri, ci
            table_id += 1


def _all_paragraphs(doc: Document) -> list[Paragraph]:
    """List all body + table cell paragraphs in document order."""
    return [para for para, _tid, _r, _c in _iter_body_blocks(doc)]


# --- Structural detection: plaintiff and defendant names from caption ---
def _detect_party_names_from_caption(para_list: list[tuple[str, Paragraph]]) -> tuple[str | None, str | None]:
    """
    In caption region: find ALL CAPS line that precedes "Plaintiff" -> plaintiff_name;
    find ALL CAPS line that precedes "Defendant" -> defendant_name.
    Returns (plaintiff_name, defendant_name); either can be None.
    """
    plaintiff_name: str | None = None
    defendant_name: str | None = None
    for i, (text, _) in enumerate(para_list):
        t = text.strip()
        if not t:
            continue
        lower = t.lower()
        next_text = ""
        if i + 1 < len(para_list):
            next_text = para_list[i + 1][0].strip().lower()
        # ALL CAPS line ending with comma, next line is "Plaintiff," or "Plaintiff"
        if t.isupper() and len(t) <= 80 and t.endswith(","):
            if next_text in ("plaintiff,", "plaintiff", "claimant,", "claimant") and plaintiff_name is None:
                plaintiff_name = t
            elif next_text in ("defendant.", "defendant", "respondent.", "respondent") and defendant_name is None:
                defendant_name = t
        # ALL CAPS line (with or without comma) when next line is "Defendant." — some docs omit comma
        if defendant_name is None and t.isupper() and len(t) <= 80 and next_text in ("defendant.", "defendant", "respondent.", "respondent"):
            defendant_name = t.rstrip(".,").strip() or t
        if plaintiff_name is None and t.isupper() and len(t) <= 80 and next_text in ("plaintiff,", "plaintiff", "claimant,", "claimant"):
            plaintiff_name = t.rstrip(".,").strip() or t
        # Same line: "NAME, Plaintiff," or "NAME, Defendant."
        if plaintiff_name is None and "plaintiff" in lower:
            m = re.match(r"^([A-Z][A-Z\s\',\.\-]+?)\s*,\s*Plaintiff[\s,\.]", t, re.I)
            if m:
                plaintiff_name = m.group(1).strip()
                if plaintiff_name.endswith(","):
                    plaintiff_name = plaintiff_name.rstrip(",").strip()
        if defendant_name is None and "defendant" in lower:
            m = re.match(r"^([A-Z][A-Z\s\',\.\-]+?)\s*,\s*Defendant[\s\.]", t, re.I)
            if m:
                defendant_name = m.group(1).strip()
                if defendant_name.endswith("."):
                    defendant_name = defendant_name.rstrip(".").strip()
    # Fallback: "Defendant." as standalone line — use previous non-empty line as defendant name
    if defendant_name is None:
        for i, (text, _) in enumerate(para_list):
            t = text.strip().lower()
            if t in ("defendant.", "defendant", "respondent.", "respondent") and i > 0:
                prev = para_list[i - 1][0].strip()
                if prev and len(prev) <= 80 and (prev.isupper() or re.match(r"^[A-Z][A-Z\s\',\.\-]+$", prev)):
                    defendant_name = prev.rstrip(".,").strip() or prev
                    break
    if plaintiff_name is None:
        for i, (text, _) in enumerate(para_list):
            t = text.strip().lower()
            if t in ("plaintiff,", "plaintiff", "claimant,", "claimant") and i > 0:
                prev = para_list[i - 1][0].strip()
                if prev and len(prev) <= 80 and (prev.isupper() or re.match(r"^[A-Z][A-Z\s\',\.\-]+$", prev)):
                    plaintiff_name = prev.rstrip(".,").strip() or prev
                    break
    return plaintiff_name, defendant_name


# --- Index No. / Date Filed value extraction ---
INDEX_NO_PATTERN = re.compile(r"Index\s+No\.?\s*:?\s*(.+)$", re.I)
DATE_FILED_PATTERN = re.compile(r"Date\s+Filed\s*:?\s*(.+)$", re.I)


def _extract_index_and_date_values(para_list: list[tuple[str, Paragraph]]) -> tuple[str | None, str | None]:
    """From first matching paragraphs, extract the value after 'Index No.:' and 'Date Filed:'."""
    index_value: str | None = None
    date_value: str | None = None
    for text, _ in para_list:
        t = text.strip()
        if index_value is None:
            m = INDEX_NO_PATTERN.match(t)
            if m:
                v = m.group(1).strip()
                if v and not v.startswith("{{"):
                    index_value = v
        if date_value is None:
            m = DATE_FILED_PATTERN.match(t)
            if m:
                v = m.group(1).strip()
                if v and not v.startswith("{{"):
                    date_value = v
        if index_value is not None and date_value is not None:
            break
    return index_value, date_value


# --- Apply replacements at paragraph level (all occurrences), count ---
def _apply_replacements_to_paragraph(
    para: Paragraph,
    replacements: dict[str, str],
    counts: dict[str, int],
) -> None:
    """
    Get full paragraph text, apply all replacements (all occurrences each), rebuild single run.
    replacements: from_text -> to_text (e.g. "JOHN DOE" -> "{{PLAINTIFF_NAME}}").
    counts: key (placeholder key or label) -> number of replacements made.
    """
    full = paragraph_full_text(para)
    if not full:
        return
    changed = False
    for from_str, to_str in replacements.items():
        if not from_str or from_str not in full:
            continue
        # Count occurrences we're replacing
        n = full.count(from_str)
        if n > 0:
            key = to_str if to_str.startswith("{{") else from_str
            # Normalize key for counting (e.g. {{PLAINTIFF_NAME}} -> PLAINTIFF_NAME)
            if key.startswith("{{") and key.endswith("}}"):
                key = key[2:-2]
            counts[key] = counts.get(key, 0) + n
            full = full.replace(from_str, to_str)
            changed = True
    if changed:
        rebuild_paragraph_single_run(para, full)


def convert_sample_to_template(
    sample_path: str | Path,
    output_path: str | Path,
) -> dict[str, int]:
    """
    Open sample DOCX, replace dynamic regions with placeholders (structural rules, ALL occurrences), save.
    Returns replacement counts per placeholder key for logging.
    """
    sample_path = Path(sample_path)
    output_path = Path(output_path)
    doc = Document(sample_path)

    para_list_with_text: list[tuple[str, Paragraph]] = []
    for para, _tid, _r, _c in _iter_body_blocks(doc):
        para_list_with_text.append((paragraph_full_text(para), para))

    counts: dict[str, int] = {}

    # 1) Detect party names and index/date values from structure
    plaintiff_name, defendant_name = _detect_party_names_from_caption(para_list_with_text)
    index_value, date_value = _extract_index_and_date_values(para_list_with_text)

    # Replace longer string first to avoid substring issues. For party names, replace BOTH
    # the exact detected form and the form without trailing comma/period so every occurrence is replaced.
    replacements: dict[str, str] = {}
    if defendant_name:
        replacements[defendant_name] = "{{DEFENDANT_NAME}}"
        # Also replace name without trailing punctuation (e.g. in body or repeated caption)
        base = defendant_name.rstrip(".,").strip()
        if base and base != defendant_name:
            replacements[base] = "{{DEFENDANT_NAME}}"
    if plaintiff_name:
        replacements[plaintiff_name] = "{{PLAINTIFF_NAME}}"
        base = plaintiff_name.rstrip(".,").strip()
        if base and base != plaintiff_name:
            replacements[base] = "{{PLAINTIFF_NAME}}"
    if index_value:
        replacements["Index No.: " + index_value] = "Index No.: {{INDEX_NO}}"
        replacements["Index No. " + index_value] = "Index No. {{INDEX_NO}}"
        replacements["INDEX NO.: " + index_value] = "INDEX NO.: {{INDEX_NO}}"
    if date_value:
        replacements["Date Filed: " + date_value] = "Date Filed: {{DATE_FILED}}"
        replacements["Date Filed " + date_value] = "Date Filed {{DATE_FILED}}"
        replacements["DATE FILED: " + date_value] = "DATE FILED: {{DATE_FILED}}"

    # Regex for Index No. / Date Filed (replace value part only; all occurrences in doc)
    index_prefix_re = re.compile(r"(Index\s+No\.?\s*:?\s*)([^\n{]+?)(?=\s*$|\s*\n)", re.I)
    date_prefix_re = re.compile(r"(Date\s+Filed\s*:?\s*)([^\n{]+?)(?=\s*$|\s*\n)", re.I)

    for para, _tid, _r, _c in _iter_body_blocks(doc):
        full = paragraph_full_text(para)
        if not full:
            continue
        # Apply string replacements (party names, literal "Index No.: value" etc.)
        _apply_replacements_to_paragraph(para, replacements, counts)
        full = paragraph_full_text(para)
        # Regex fallback for Index No. / Date Filed (any remaining)
        if re.search(r"Index\s+No\.?\s*:?\s*[^{]", full, re.I):
            new_full, n = index_prefix_re.subn(r"\1{{INDEX_NO}}", full)
            if n > 0:
                counts["INDEX_NO"] = counts.get("INDEX_NO", 0) + n
                rebuild_paragraph_single_run(para, new_full)
                full = new_full
        if re.search(r"Date\s+Filed\s*:?\s*[^{]", full, re.I):
            new_full, n = date_prefix_re.subn(r"\1{{DATE_FILED}}", full)
            if n > 0:
                counts["DATE_FILED"] = counts.get("DATE_FILED", 0) + n
                rebuild_paragraph_single_run(para, new_full)

    # 2) Venue / plaintiff residence (first matching paragraph only)
    for para, _tid, _r, _c in _iter_body_blocks(doc):
        full = paragraph_full_text(para).strip().lower()
        if ("basis of venue" in full or "venue is" in full or "plaintiff's residence" in full) and "{{" not in paragraph_full_text(para):
            rebuild_paragraph_single_run(para, "{{PLAINTIFF_RESIDENCE}}")
            counts["PLAINTIFF_RESIDENCE"] = counts.get("PLAINTIFF_RESIDENCE", 0) + 1
            break

    # 3) Numbered allegations block -> single placeholder
    allegation_start = re.compile(r"^\s*(?:\d+\.\s+|That\s+|By\s+reason\s+of|Pursuant\s+to)", re.I)
    cause_title = re.compile(r"^(?:AS\s+AND\s+FOR\s+[A-Z]+\s+CAUSE\s+OF\s+ACTION\s*:?\s*)?([A-Z][A-Z\s]+)\s*$")
    block_marker = "{{CAUSE_OF_ACTION_1_PARAGRAPHS__BLOCK}}"
    paras = _all_paragraphs(doc)
    block_start_idx: int | None = None
    block_end_idx: int | None = None
    for i, para in enumerate(paras):
        text = paragraph_full_text(para).strip()
        if cause_title.match(text) and len(text) < 100:
            j = i + 1
            while j < len(paras):
                ptext = paragraph_full_text(paras[j]).strip()
                if not ptext:
                    j += 1
                    continue
                if allegation_start.match(ptext) or (re.match(r"^\d+\.\s+", ptext) and len(ptext) > 20):
                    if block_start_idx is None:
                        block_start_idx = j
                    block_end_idx = j
                    j += 1
                elif "cause of action" in ptext.lower() or "wherefore" in ptext.lower():
                    break
                else:
                    j += 1
            if block_start_idx is not None and block_end_idx is not None:
                to_replace = [paras[k] for k in range(block_start_idx, block_end_idx + 1)]
                rebuild_paragraph_single_run(to_replace[0], block_marker)
                for p in to_replace[1:]:
                    parent = p._element.getparent()
                    if parent is not None:
                        parent.remove(p._element)
                counts["CAUSE_OF_ACTION_1_PARAGRAPHS__BLOCK"] = 1
            break

    # 4) Signature block heuristics (single-run replacements)
    phone_re = re.compile(r"^\(?\d{3}\)?\s*\d{3}[-.\s]?\d{4}\s*$")
    for para, _tid, _r, _c in _iter_body_blocks(doc):
        full = paragraph_full_text(para).strip()
        if not full or "{{" in full:
            continue
        lower = full.lower()
        if ", esq." in lower or ", esq" in lower:
            if len(full) < 60:
                rebuild_paragraph_single_run(para, "{{ATTORNEY_NAME}}")
                counts["ATTORNEY_NAME"] = counts.get("ATTORNEY_NAME", 0) + 1
        elif phone_re.match(full):
            rebuild_paragraph_single_run(para, "{{PHONE}}")
            counts["PHONE"] = counts.get("PHONE", 0) + 1
        elif ("pllc" in lower or "llc" in lower or "p.c." in lower) and "law" in lower and len(full) <= 60:
            rebuild_paragraph_single_run(para, "{{FIRM_NAME}}")
            counts["FIRM_NAME"] = counts.get("FIRM_NAME", 0) + 1
        elif re.match(r"^\d+[\w\s,]+(?:Avenue|Street|Boulevard|Road|Drive|Lane|Suite|Floor)", full, re.I):
            rebuild_paragraph_single_run(para, "{{FIRM_ADDRESS}}")
            counts["FIRM_ADDRESS"] = counts.get("FIRM_ADDRESS", 0) + 1
        elif re.match(r"^[A-Za-z\s]+,?\s*(?:New York|NY|Connecticut|CT)\s+\d{5}", full, re.I):
            rebuild_paragraph_single_run(para, "{{FIRM_ADDRESS}}")
            counts["FIRM_ADDRESS"] = counts.get("FIRM_ADDRESS", 0) + 1
        elif "dated:" in lower and len(full) < 80:
            rebuild_paragraph_single_run(para, "{{SIGNATURE_DATE}}")
            counts["SIGNATURE_DATE"] = counts.get("SIGNATURE_DATE", 0) + 1

    # Log counts
    log.info("Template build replacement counts: %s", counts)
    for key, n in sorted(counts.items()):
        log.info("  %s: %d", key, n)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return counts
