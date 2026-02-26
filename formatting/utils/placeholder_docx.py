"""Fixed-template document generation: open template, replace placeholders, save.

Run-safe: placeholders and names can span multiple runs. We replace at the PARAGRAPH level:
get full paragraph text, apply replacements, rebuild paragraph into a SINGLE run.
Paragraph formatting (style, alignment, tab stops, borders) is never changed.
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def paragraph_full_text(paragraph) -> str:
    """Full text of paragraph (all runs concatenated). Use this instead of paragraph.text when placeholders span runs."""
    return "".join(r.text or "" for r in paragraph.runs)


def rebuild_paragraph_single_run(paragraph, new_text: str) -> None:
    """
    Replace paragraph content with new_text in a single run. Preserves paragraph formatting
    (style, alignment, tab stops, borders). Removes all existing runs and adds one run.
    """
    for r in list(paragraph.runs)[::-1]:
        r._element.getparent().remove(r._element)
    run = paragraph.add_run(new_text)
    run.bold = False
    run.italic = False
    run.underline = False


def safe_replace_paragraph(paragraph, old: str, new: str, case_insensitive: bool = False) -> bool:
    """
    Run-safe: replace old with new in this paragraph only. Uses full paragraph text;
    if old not in full text, return False. Otherwise rebuild single run with new_text.
    Preserves paragraph formatting. Returns True if replacement was made.
    If case_insensitive, the first occurrence of old (ignoring case) is replaced with new.
    """
    full_text = paragraph_full_text(paragraph)
    if not old:
        return False
    if case_insensitive:
        lower_full = full_text.lower()
        old_lower = old.lower()
        if old_lower not in lower_full:
            return False
        # Replace every occurrence (case-insensitive); build new_text by repeated replace
        new_text = full_text
        start = 0
        while True:
            idx = new_text.lower().find(old_lower, start)
            if idx < 0:
                break
            new_text = new_text[:idx] + new + new_text[idx + len(old):]
            start = idx + len(new)
    else:
        if old not in full_text:
            return False
        new_text = full_text.replace(old, new)
    rebuild_paragraph_single_run(paragraph, new_text)
    return True


def replace_globally(doc: Document, old: str, new: str, case_insensitive: bool = False) -> int:
    """
    Replace old with new in every paragraph (body + table cells). Run-safe at paragraph level.
    If case_insensitive, replaces the first occurrence of old (ignoring case) in each paragraph.
    Returns number of paragraphs in which a replacement was made.
    """
    if not old:
        return 0
    count = 0
    for p in _iter_all_paragraphs(doc):
        if safe_replace_paragraph(p, old, new, case_insensitive=case_insensitive):
            count += 1
    return count


def _iter_all_paragraphs(doc: Document):
    """Yield every paragraph in document order (body + table cells)."""
    body = doc.element.body
    qp, qt = qn("w:p"), qn("w:tbl")
    for child in body.iterchildren():
        if child.tag == qp:
            yield Paragraph(child, doc)
        elif child.tag == qt:
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para


def iter_all_paragraphs_including_headers_footers(doc: Document):
    """
    Yield every paragraph in the document: body (including table cells), then each section's
    header paragraphs and table cells, then each section's footer paragraphs and table cells.
    Use for template-wide replace and validation.
    """
    for p in _iter_all_paragraphs(doc):
        yield p
    for section in doc.sections:
        try:
            header = section.header
            if header is not None:
                for p in header.paragraphs:
                    yield p
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p
        except Exception:
            pass
        try:
            footer = section.footer
            if footer is not None:
                for p in footer.paragraphs:
                    yield p
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p
        except Exception:
            pass


def replace_globally_full(doc: Document, old: str, new: str, case_insensitive: bool = False) -> int:
    """
    Replace old with new in every paragraph: body, table cells, headers, footers.
    Run-safe at paragraph level. Returns number of paragraphs in which a replacement was made.
    """
    if not old:
        return 0
    count = 0
    for p in iter_all_paragraphs_including_headers_footers(doc):
        if safe_replace_paragraph(p, old, new, case_insensitive=case_insensitive):
            count += 1
    return count


def document_contains_value(doc: Document, value: str) -> bool:
    """Return True if value appears in any paragraph (body, tables, headers, footers)."""
    if not value:
        return False
    for p in iter_all_paragraphs_including_headers_footers(doc):
        if value in paragraph_full_text(p):
            return True
    return False


def replace_marker_across_runs(paragraph, marker: str, value: str) -> bool:
    """
    Replace a placeholder marker that may be split across runs (e.g. Run1 '{{LIC' Run2 'ENSE_PLATE}}').
    Concatenates run texts, finds marker, then rewrites the first involved run with the
    replaced text and clears the others. Returns True if replacement was made.
    """
    full = "".join(r.text or "" for r in paragraph.runs)
    idx = full.find(marker)
    if idx < 0:
        return False
    end = idx + len(marker)

    pos = 0
    run_spans = []
    for ri, r in enumerate(paragraph.runs):
        run_len = len(r.text or "")
        run_spans.append((ri, pos, pos + run_len))
        pos += run_len

    involved = [s for s in run_spans if not (s[2] <= idx or s[1] >= end)]
    if not involved:
        return False

    first_i = involved[0][0]
    last_i = involved[-1][0]

    new_full = full[:idx] + value + full[end:]

    paragraph.runs[first_i].text = new_full
    for ri in range(first_i + 1, len(paragraph.runs)):
        if ri <= last_i:
            paragraph.runs[ri].text = ""
    return True


def _replace_placeholder_in_paragraph(p, replacements: dict[str, str]) -> None:
    """Replace placeholders in one paragraph. Run-safe: use paragraph_full_text, apply all replacements, rebuild single run."""
    full_text = paragraph_full_text(p)
    changed = False
    for marker, value in replacements.items():
        if marker in full_text:
            full_text = full_text.replace(marker, value)
            changed = True
    if changed:
        rebuild_paragraph_single_run(p, full_text)


def replace_placeholders(doc: Document, replacements: dict[str, str]) -> None:
    """Replace {{KEY}} placeholders in all paragraphs and table cells. Run-safe:
    when a placeholder is found in paragraph text, the paragraph is rebuilt with
    one clean run to avoid partial replacement and style leakage."""
    for p in doc.paragraphs:
        _replace_placeholder_in_paragraph(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_placeholder_in_paragraph(p, replacements)


def generate_from_template(
    template_path: str | Path,
    replacements: dict[str, str],
    output_path: str | Path | None = None,
) -> Document:
    """
    Load a DOCX template, replace all placeholders, optionally save to file.
    Returns the modified Document so you can add dynamic sections or save elsewhere.
    """
    doc = Document(template_path)
    replace_placeholders(doc, replacements)
    if output_path is not None:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
    return doc
