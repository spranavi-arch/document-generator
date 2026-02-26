"""
Sample-to-template pipeline: convert sample.docx into template.docx with placeholders.

Phase 1: Run map extraction, heuristic field detection, run-safe placeholder insertion.
Phase 2: Schema generation from placeholders.

Entry: sample_to_template(sample_path, template_path, schema_path)
"""

from __future__ import annotations

import json
import re
from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

try:
    from utils.style_extractor import iter_body_blocks
except ImportError:
    from style_extractor import iter_body_blocks


# ---------------------------------------------------------------------------
# Run map (Phase 1.1)
# ---------------------------------------------------------------------------


def _get_paragraph_border_info(para) -> bool:
    """True if paragraph has a bottom border (divider line) — do not replace."""
    try:
        p = para._p
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            return False
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            return False
        return pBdr.find(qn("w:bottom")) is not None
    except Exception:
        return False


def extract_run_map(doc: Document) -> list[dict[str, Any]]:
    """
    Extract all text runs with positions in document order.
    Returns list of run records: para_id, run_index, run_text, style_name, table_id, row, col,
    para_ref (paragraph object), run_ref (run object), has_bottom_border.
    """
    run_map: list[dict[str, Any]] = []
    para_id = 0
    for para, table_id, row, col in iter_body_blocks(doc):
        try:
            style_name = para.style.name if para.style else None
        except Exception:
            style_name = None
        full_text = (para.text or "").replace("\n", " ")
        has_border = _get_paragraph_border_info(para)
        for run_index, run in enumerate(para.runs):
            run_text = (run.text or "")
            run_map.append({
                "para_id": para_id,
                "run_index": run_index,
                "run_text": run_text,
                "full_para_text": full_text,
                "style_name": style_name,
                "table_id": table_id,
                "row": row,
                "col": col,
                "para_ref": para,
                "run_ref": run,
                "has_bottom_border": has_border,
            })
        para_id += 1
    return run_map


def build_paragraph_map(doc: Document) -> list[dict[str, Any]]:
    """Build a list of paragraph records (para_id, paragraph, full_text, style_name, table_id, row, col) in document order."""
    para_map: list[dict[str, Any]] = []
    for para_id, (para, table_id, row, col) in enumerate(iter_body_blocks(doc)):
        full_text = (para.text or "").strip()
        try:
            style_name = para.style.name if para.style else None
        except Exception:
            style_name = None
        para_map.append({
            "para_id": para_id,
            "paragraph": para,
            "full_text": full_text,
            "style_name": style_name,
            "table_id": table_id,
            "row": row,
            "col": col,
        })
    return para_map


# ---------------------------------------------------------------------------
# Heuristic field detection (Phase 1.2)
# ---------------------------------------------------------------------------

PHONE_RE = re.compile(r"^\(?\d{3}\)?\s*\d{3}[-.\s]?\d{4}\s*$")
INDEX_NO_RE = re.compile(r"^\s*Index\s+No\.?\s*:?\s*(.+)$", re.I)
DATE_FILED_RE = re.compile(r"^\s*Date\s+Filed\s*:?\s*(.+)$", re.I)
VENUE_RE = re.compile(r"basis\s+of\s+venue|venue\s+is", re.I)
WHEREFORE_RE = re.compile(r"^\s*WHEREFORE\s*,?", re.I)
CAUSE_TITLE_RE = re.compile(r"^(?:AS\s+AND\s+FOR\s+[A-Z]+\s+CAUSE\s+OF\s+ACTION\s*:?\s*)?([A-Z][A-Z\s]+)\s*$")
ALLEGATION_START_RE = re.compile(r"^\s*(?:That\s+|By\s+reason\s+of|Pursuant\s+to)", re.I)


def _detect_scalar_candidates(para_map: list[dict]) -> list[dict[str, Any]]:
    """Detect placeholder candidates from paragraph text. Returns list of { placeholder_key, para_id, start_char, end_char }."""
    candidates: list[dict[str, Any]] = []
    seen_plaintiff = False
    seen_defendant = False
    in_caption = True
    for rec in para_map:
        para_id = rec["para_id"]
        text = rec["full_text"]
        if not text:
            continue
        t = text.strip()
        lower = t.lower()
        # Skip divider-only paragraphs
        if len(t) < 3:
            continue

        # Index No.: value
        m = INDEX_NO_RE.match(t)
        if m:
            value = m.group(1).strip()
            if value and not value.startswith("{{"):
                # Replace the value part (after "Index No.: ")
                prefix = "Index No.: "
                if prefix in t:
                    start = t.index(prefix) + len(prefix)
                    end = len(t)
                    candidates.append({"placeholder_key": "INDEX_NO", "para_id": para_id, "start_char": start, "end_char": end})
            continue

        # Date Filed: value
        m = DATE_FILED_RE.match(t)
        if m:
            value = m.group(1).strip()
            if value and not value.startswith("{{"):
                prefix = "Date Filed:" if "Date Filed:" in t else "Date Filed"
                if prefix in t:
                    start = t.index(prefix) + len(prefix)
                    end = len(t)
                    candidates.append({"placeholder_key": "DATE_FILED", "para_id": para_id, "start_char": start, "end_char": end})
            continue

        # Phone
        if PHONE_RE.match(t):
            candidates.append({"placeholder_key": "SIGNATURE_BLOCK.PHONE", "para_id": para_id, "start_char": 0, "end_char": len(t)})
            continue

        # Venue basis
        if VENUE_RE.search(t) and len(t) < 300:
            candidates.append({"placeholder_key": "VENUE_BASIS", "para_id": para_id, "start_char": 0, "end_char": len(t)})
            continue

        # WHEREFORE
        if WHEREFORE_RE.match(t) or (lower.startswith("wherefore") and len(t) < 500):
            candidates.append({"placeholder_key": "WHEREFORE", "para_id": para_id, "start_char": 0, "end_char": len(t)})
            in_caption = False
            continue

        # Body starts after caption; mark end of caption for party names
        if "to the above named defendant" in lower or "you are hereby summoned" in lower:
            in_caption = False
        if "attorneys for plaintiff" in lower or "attorney for plaintiff" in lower:
            in_caption = False

        # ALL CAPS name line in caption: check next paragraph for "Plaintiff," / "Defendant."
        if in_caption and len(t) <= 80 and t.isupper() and t.endswith(","):
            next_text = ""
            if para_id + 1 < len(para_map):
                next_text = (para_map[para_id + 1]["full_text"] or "").strip().lower()
            if next_text in ("plaintiff,", "plaintiff", "claimant,", "claimant") and not seen_plaintiff:
                seen_plaintiff = True
                candidates.append({"placeholder_key": "PLAINTIFF_NAME", "para_id": para_id, "start_char": 0, "end_char": len(t)})
                continue
            if next_text in ("defendant.", "defendant", "respondent.", "respondent") and not seen_defendant:
                seen_defendant = True
                candidates.append({"placeholder_key": "DEFENDANT_NAME", "para_id": para_id, "start_char": 0, "end_char": len(t)})
                continue
            # Fallback: line contains "plaintiff" or "defendant" (e.g. "PLAINTIFF," as the line itself)
            if not seen_plaintiff and "plaintiff" in lower:
                seen_plaintiff = True
                candidates.append({"placeholder_key": "PLAINTIFF_NAME", "para_id": para_id, "start_char": 0, "end_char": len(t)})
                continue
            if not seen_defendant and "defendant" in lower:
                seen_defendant = True
                candidates.append({"placeholder_key": "DEFENDANT_NAME", "para_id": para_id, "start_char": 0, "end_char": len(t)})
                continue

        # Address-like lines (street, city state zip)
        if re.match(r"^\d+[\w\s,]+(Avenue|Street|Boulevard|Road|Drive|Lane|Suite|Floor)", t, re.I):
            candidates.append({"placeholder_key": "SIGNATURE_BLOCK.ADDRESS_LINE_1", "para_id": para_id, "start_char": 0, "end_char": len(t)})
            continue
        if re.match(r"^[A-Za-z\s]+,?\s*(New York|NY|Connecticut|CT)\s+\d{5}", t, re.I):
            candidates.append({"placeholder_key": "SIGNATURE_BLOCK.ADDRESS_LINE_2", "para_id": para_id, "start_char": 0, "end_char": len(t)})
            continue

        # Firm name (often all caps, before "Attorneys for Plaintiff")
        if "attorneys for plaintiff" in lower or "attorney for plaintiff" in lower:
            continue
        if in_caption is False and len(t) <= 60 and (t.isupper() or "pllc" in lower or "llc" in lower) and "law" in lower:
            candidates.append({"placeholder_key": "SIGNATURE_BLOCK.FIRM", "para_id": para_id, "start_char": 0, "end_char": len(t)})
            continue

        # Attorney name (e.g. "MICHAEL COHAN, ESQ.")
        if ", esq." in lower or ", esq" in lower:
            if len(t) < 60:
                candidates.append({"placeholder_key": "SIGNATURE_BLOCK.ATTORNEY", "para_id": para_id, "start_char": 0, "end_char": len(t)})
            continue

    return candidates


def _detect_block_candidates(para_map: list[dict]) -> list[dict[str, Any]]:
    """Detect block placeholders (e.g. cause of action paragraphs). Returns list of { placeholder_key, para_id } (whole paragraph)."""
    block_candidates: list[dict[str, Any]] = []
    i = 0
    while i < len(para_map):
        rec = para_map[i]
        text = (rec["full_text"] or "").strip()
        lower = text.lower()
        # Cause of action title (e.g. "NEGLIGENCE" or "AS AND FOR A FIRST CAUSE OF ACTION: NEGLIGENCE")
        if CAUSE_TITLE_RE.match(text) and len(text) < 100:
            block_candidates.append({"placeholder_key": "CAUSE_OF_ACTION_1_TITLE", "para_id": rec["para_id"], "start_char": 0, "end_char": len(text)})
            i += 1
            # Next paragraphs that look like allegations → one block placeholder
            allegation_paras = []
            j = i
            while j < len(para_map):
                ptext = (para_map[j]["full_text"] or "").strip()
                if not ptext:
                    j += 1
                    continue
                if ALLEGATION_START_RE.match(ptext) or (re.match(r"^\d+\.\s+", ptext) and len(ptext) > 20):
                    allegation_paras.append(para_map[j]["para_id"])
                    j += 1
                elif "cause of action" in ptext.lower() or "wherefore" in ptext.lower():
                    break
                else:
                    j += 1
            if allegation_paras:
                block_candidates.append({
                    "placeholder_key": "CAUSE_OF_ACTION_1_PARAS__BLOCK",
                    "para_id": allegation_paras[0],
                    "block_para_ids": allegation_paras,
                    "is_block": True,
                })
                i = j
                continue
        i += 1
    return block_candidates


# ---------------------------------------------------------------------------
# Optional LLM for ambiguous fields (Phase 1.3)
# ---------------------------------------------------------------------------


def get_ambiguous_regions(
    para_map: list[dict],
    scalar_candidates: list[dict],
    block_candidates: list[dict],
    max_value_len: int = 200,
) -> list[dict[str, Any]]:
    """
    Return regions that heuristics did not classify. Each region has para_id, start_char, end_char,
    value (text), context_before (list of up to 3 paragraph texts), context_after.
    Used for optional LLM classification.
    """
    covered_paras: set[int] = set()
    for c in block_candidates:
        covered_paras.add(c["para_id"])
    covered_spans: set[tuple[int, int, int]] = set()
    for c in scalar_candidates:
        covered_spans.add((c["para_id"], c.get("start_char", 0), c.get("end_char", 0)))
    regions: list[dict[str, Any]] = []
    for rec in para_map:
        para_id = rec["para_id"]
        text = rec["full_text"]
        if not text or len(text.strip()) < 2:
            continue
        if para_id in covered_paras:
            continue
        t = text.strip()
        if any(para_id == p and start <= 0 and end >= len(t) for p, start, end in covered_spans):
            continue
        value = t[:max_value_len]
        context_before = [para_map[i]["full_text"].strip()[:80] for i in range(max(0, para_id - 3), para_id) if para_map[i]["full_text"].strip()]
        context_after = [para_map[i]["full_text"].strip()[:80] for i in range(para_id + 1, min(len(para_map), para_id + 4)) if para_map[i]["full_text"].strip()]
        regions.append({
            "para_id": para_id,
            "start_char": 0,
            "end_char": len(t),
            "value": value,
            "context_before": context_before,
            "context_after": context_after,
        })
    return regions


def build_classification_prompt(region: dict[str, Any]) -> str:
    """Build prompt for LLM: what placeholder should replace this value? Return JSON: {\"placeholder\":\"...\", \"confidence\":0.0-1.0}."""
    before = "\n".join(region.get("context_before") or [])
    after = "\n".join(region.get("context_after") or [])
    value = region.get("value") or ""
    return (
        "Context before:\n" + (before or "(none)") + "\n\n"
        "Value to classify:\n" + value + "\n\n"
        "Context after:\n" + (after or "(none)") + "\n\n"
        "What placeholder key should replace this value? Return only JSON: {\"placeholder\": \"KEY_NAME\", \"confidence\": 0.0-1.0}."
    )


def parse_classification_response(raw: str) -> dict[str, Any] | None:
    """Parse LLM response to classification prompt. Returns {\"placeholder\": \"...\", \"confidence\": float} or None."""
    text = (raw or "").strip()
    if not text:
        return None
    # Strip code fence if present
    if text.startswith("```"):
        lines = text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        text = "\n".join(lines)
    try:
        data = json.loads(text)
        if isinstance(data, dict) and "placeholder" in data:
            conf = data.get("confidence")
            if conf is not None:
                data["confidence"] = float(conf)
            return data
    except (json.JSONDecodeError, TypeError, ValueError):
        pass
    return None


# ---------------------------------------------------------------------------
# Run-safe placeholder insertion (Phase 1.4)
# ---------------------------------------------------------------------------


def _rebuild_paragraph_with_placeholders(paragraph, full_text: str, segments: list[tuple[int, int, str | None]]) -> None:
    """
    Rebuild paragraph so that segments are replaced. segments: list of (start, end, placeholder_key).
    placeholder_key is None for unchanged text. Sort by start; non-overlapping.
    """
    parts: list[tuple[int, int, str]] = []
    pos = 0
    for start, end, key in sorted(segments, key=lambda x: x[0]):
        if start > pos:
            parts.append((pos, start, full_text[pos:start]))
        if key:
            parts.append((start, end, "{{" + key + "}}"))
        pos = end
    if pos < len(full_text):
        parts.append((pos, len(full_text), full_text[pos:]))
    for r in list(paragraph.runs)[::-1]:
        r._element.getparent().remove(r._element)
    for _s, _e, repl_text in parts:
        if repl_text:
            run = paragraph.add_run(repl_text)
            run.bold = False
            run.italic = False
            run.underline = False


def _apply_scalar_candidates(doc: Document, para_map: list[dict], candidates: list[dict]) -> None:
    """Apply scalar replacements: group by para_id, then rebuild each paragraph with placeholder runs."""
    by_para: dict[int, list[dict]] = defaultdict(list)
    for c in candidates:
        if c.get("is_block"):
            continue
        by_para[c["para_id"]].append(c)
    for para_id in sorted(by_para.keys()):
        if para_id >= len(para_map):
            continue
        list_c = by_para[para_id]
        rec = para_map[para_id]
        para = rec["paragraph"]
        full_text = rec["full_text"] or ""
        if not full_text:
            full_text = " "
        segments: list[tuple[int, int, str | None]] = []
        for c in sorted(list_c, key=lambda x: x["start_char"]):
            start = c["start_char"]
            end = c.get("end_char", len(full_text))
            segments.append((start, end, c["placeholder_key"]))
        merged: list[tuple[int, int, str | None]] = []
        pos = 0
        for start, end, key in segments:
            if start > pos:
                merged.append((pos, start, None))
            merged.append((start, end, key))
            pos = end
        if pos < len(full_text):
            merged.append((pos, len(full_text), None))
        _rebuild_paragraph_with_placeholders(para, full_text, merged)


def _apply_block_candidates(doc: Document, para_map: list[dict], block_candidates: list[dict]) -> None:
    """Replace block paragraphs with a single paragraph containing {{KEY__BLOCK}}."""
    for c in block_candidates:
        if not c.get("is_block"):
            continue
        key = c["placeholder_key"]
        para_ids = c.get("block_para_ids") or [c["para_id"]]
        if not para_ids:
            continue
        first_id = para_ids[0]
        if first_id >= len(para_map):
            continue
        rec = para_map[first_id]
        para = rec["paragraph"]
        for r in list(para.runs)[::-1]:
            r._element.getparent().remove(r._element)
        run = para.add_run("{{" + key + "}}")
        run.bold = False
        run.italic = False
        run.underline = False
        for pid in para_ids[1:]:
            if pid >= len(para_map):
                continue
            other = para_map[pid]["paragraph"]
            for r in list(other.runs)[::-1]:
                r._element.getparent().remove(r._element)
            other.add_run("")


# ---------------------------------------------------------------------------
# Schema generation (Phase 2)
# ---------------------------------------------------------------------------

def detect_placeholder_keys(doc: Document, doc_type: str = "SummonsAndComplaint") -> list[str]:
    """
    Run heuristic detection on a document without modifying it. Returns sorted list of
    placeholder keys that would be used. Use for a secondary sample to merge keys.
    """
    para_map = build_paragraph_map(doc)
    scalar_candidates = _detect_scalar_candidates(para_map)
    block_candidates = _detect_block_candidates(para_map)
    keys = set()
    for c in scalar_candidates:
        keys.add(c["placeholder_key"])
    for c in block_candidates:
        keys.add(c["placeholder_key"])
    return sorted(keys)


# Document-type-specific placeholder sets: which keys are typically required for that doc type.
DOC_TYPE_PLACEHOLDERS: dict[str, set[str]] = {
    "SummonsAndComplaint": {
        "INDEX_NO", "DATE_FILED", "PLAINTIFF_NAME", "DEFENDANT_NAME",
        "VENUE_BASIS", "WHEREFORE", "CAUSE_OF_ACTION_1_TITLE", "CAUSE_OF_ACTION_1_PARAS__BLOCK",
    },
    "Motion": {"INDEX_NO", "DATE_FILED", "PLAINTIFF_NAME", "DEFENDANT_NAME"},
}


def build_schema_from_placeholders(placeholder_keys: list[str], doc_type: str = "SummonsAndComplaint") -> dict[str, Any]:
    """Build schema.json structure from list of placeholder keys. Uses DOC_TYPE_PLACEHOLDERS when set."""
    placeholders: dict[str, dict] = {}
    required_for_type = DOC_TYPE_PLACEHOLDERS.get(doc_type) or set()
    for key in placeholder_keys:
        if key.endswith("__BLOCK"):
            placeholders[key] = {"type": "string_list", "required": True, "render": "paragraphs"}
        elif key.startswith("SIGNATURE_BLOCK."):
            placeholders[key] = {"type": "string", "required": False}
        elif key in ("INDEX_NO", "PLAINTIFF_NAME", "DEFENDANT_NAME", "VENUE_BASIS", "WHEREFORE", "CAUSE_OF_ACTION_1_TITLE"):
            placeholders[key] = {"type": "string", "required": True}
        elif key == "DATE_FILED":
            placeholders[key] = {"type": "date", "required": True}
        else:
            placeholders[key] = {"type": "string", "required": key in required_for_type}
    return {"doc_type": doc_type, "placeholders": placeholders}


def save_schema(schema: dict, schema_path: str | Path) -> None:
    """Write schema to JSON file."""
    path = Path(schema_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(schema, f, indent=2, ensure_ascii=False)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def _spans_overlap(a: dict, b: dict) -> bool:
    """True if two candidates overlap (same para_id, character ranges overlap)."""
    if a.get("para_id") != b.get("para_id"):
        return False
    s1, e1 = a.get("start_char", 0), a.get("end_char", 0)
    s2, e2 = b.get("start_char", 0), b.get("end_char", 0)
    return s1 < e2 and s2 < e1


def merge_diff_with_heuristic_candidates(
    para_map: list[dict],
    scalar_candidates: list[dict],
    block_candidates: list[dict],
    diff_candidates: list[dict],
) -> tuple[list[dict], list[dict]]:
    """
    Prefer diff-based candidates over heuristics for overlapping spans.
    Remove any scalar candidate that overlaps a diff candidate; then add all diff candidates.
    Block candidates are unchanged.
    """
    scalar_filtered = [
        c for c in scalar_candidates
        if not any(_spans_overlap(c, d) for d in diff_candidates)
    ]
    scalar_merged = scalar_filtered + list(diff_candidates)
    return scalar_merged, block_candidates


def sample_to_template(
    sample_path: str | Path,
    template_path: str | Path,
    schema_path: str | Path,
    doc_type: str = "SummonsAndComplaint",
    *,
    llm_classifier: Any = None,
    classification_confidence_threshold: float = 0.8,
) -> tuple[list[str], dict]:
    """
    Convert sample.docx to template.docx with placeholders and write schema.json.

    If llm_classifier is provided (callable taking prompt str, returning str), ambiguous
    regions are classified and merged when confidence >= classification_confidence_threshold.

    Returns (list of placeholder keys found, schema dict).
    """
    sample_path = Path(sample_path)
    template_path = Path(template_path)
    schema_path = Path(schema_path)
    doc = Document(sample_path)
    para_map = build_paragraph_map(doc)
    scalar_candidates = _detect_scalar_candidates(para_map)
    block_candidates = _detect_block_candidates(para_map)
    if callable(llm_classifier):
        ambiguous = get_ambiguous_regions(para_map, scalar_candidates, block_candidates)
        for region in ambiguous:
            prompt = build_classification_prompt(region)
            try:
                response = llm_classifier(prompt)
                parsed = parse_classification_response(response)
                if parsed and parsed.get("confidence", 0) >= classification_confidence_threshold:
                    key = (parsed.get("placeholder") or "").strip()
                    if key:
                        scalar_candidates.append({
                            "placeholder_key": key,
                            "para_id": region["para_id"],
                            "start_char": region["start_char"],
                            "end_char": region["end_char"],
                        })
            except Exception:
                pass
    _apply_scalar_candidates(doc, para_map, scalar_candidates)
    _apply_block_candidates(doc, para_map, block_candidates)
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(template_path)
    all_keys = list({c["placeholder_key"] for c in scalar_candidates + block_candidates})
    schema = build_schema_from_placeholders(all_keys, doc_type=doc_type)
    save_schema(schema, schema_path)
    return all_keys, schema


def sample_to_template_with_diff(
    primary_path: str | Path,
    template_path: str | Path,
    schema_path: str | Path,
    doc_type: str = "SummonsAndComplaint",
    diff_candidates: list[dict] | None = None,
    *,
    llm_classifier: Any = None,
    classification_confidence_threshold: float = 0.8,
) -> tuple[list[str], dict]:
    """
    Build template from primary sample, using diff-based candidates (from comparing
    primary vs secondary) when provided. Structural diff reveals variable spans;
    LLM classifies them. Heuristic detection still runs on primary; overlapping
    spans are replaced by diff-based placeholders.

    diff_candidates: list of {para_id, start_char, end_char, placeholder_key} from
    get_diff_based_candidates(primary_doc, secondary_doc, llm_callable).

    Returns (list of placeholder keys, schema dict).
    """
    primary_path = Path(primary_path)
    template_path = Path(template_path)
    schema_path = Path(schema_path)
    doc = Document(primary_path)
    para_map = build_paragraph_map(doc)
    scalar_candidates = _detect_scalar_candidates(para_map)
    block_candidates = _detect_block_candidates(para_map)
    if diff_candidates:
        scalar_candidates, block_candidates = merge_diff_with_heuristic_candidates(
            para_map, scalar_candidates, block_candidates, diff_candidates
        )
    if callable(llm_classifier):
        ambiguous = get_ambiguous_regions(para_map, scalar_candidates, block_candidates)
        for region in ambiguous:
            prompt = build_classification_prompt(region)
            try:
                response = llm_classifier(prompt)
                parsed = parse_classification_response(response)
                if parsed and parsed.get("confidence", 0) >= classification_confidence_threshold:
                    key = (parsed.get("placeholder") or "").strip()
                    if key:
                        scalar_candidates.append({
                            "placeholder_key": key,
                            "para_id": region["para_id"],
                            "start_char": region["start_char"],
                            "end_char": region["end_char"],
                        })
            except Exception:
                pass
    _apply_scalar_candidates(doc, para_map, scalar_candidates)
    _apply_block_candidates(doc, para_map, block_candidates)
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(template_path)
    all_keys = list({c["placeholder_key"] for c in scalar_candidates + block_candidates})
    schema = build_schema_from_placeholders(all_keys, doc_type=doc_type)
    save_schema(schema, schema_path)
    return all_keys, schema
